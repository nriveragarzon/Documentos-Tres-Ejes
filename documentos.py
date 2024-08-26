# Liberias 
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.section import CT_SectPr
from docx.table import _Row

def verif_ejes(session, params):
    """
    Función para verificar la existencia de datos en diferentes categorías (exportaciones, inversión y turismo)
    agrupados por diferentes criterios (CONTINENTES, HUBS, TLCS, PAISES, DEPARTAMENTOS). La función ejecuta
    consultas SQL en una sesión de Snowflake y agrega los resultados a un diccionario indicando si hay datos
    disponibles o no en cada categoría y periodo (cerrado o corrido).

    Parámetros:
    - session: Sesión activa de Snowflake.
    - params: Diccionario con los parámetros necesarios para ejecutar las consultas, incluyendo AGRUPACION,
              UNIDAD, UMBRAL, PAISES_INVERSION, PAISES_TURISMO_COD, y UNIDAD_COD.

    Retorna:
    - dict_verif: Diccionario con los resultados de la verificación, indicando si hay datos disponibles o no
                  para cada categoría y periodo.
    """

    # 1. Obtener los parámetros según sea la agrupación y unidad
    # Parámetros para los datos de exportaciones
    AGRUPACION = params['AGRUPACION']
    UNIDAD = params['UNIDAD'][0]
    UMBRAL = params['UMBRAL'][0]

    # Parámetros para los datos de inversión
    if AGRUPACION in ['CONTINENTES', 'HUBS', 'TLCS', 'PAISES']:
        PAISES_INVERSION = [pais for pais in params['PAISES_INVERSION'] if pais is not None]
        PAISES_INVERSION_sql = ', '.join(f"'{pais}'" for pais in PAISES_INVERSION)

    # Parámetros para los datos de turismo
    if AGRUPACION in ['CONTINENTES', 'HUBS', 'TLCS', 'PAISES']:
        PAISES_TURISMO = [pais for pais in params['PAISES_TURISMO_COD'] if pais is not None]
        PAISES_TURISMO_sql = ', '.join(f"'{pais}'" for pais in PAISES_TURISMO)
    if AGRUPACION in ['DEPARTAMENTOS']:
        DEPARTAMENTOS_TURISMO = [departamento for departamento in params['UNIDAD_COD'] if departamento is not None]
        DEPARTAMENTOS_TURISMO_sql = ', '.join(f"'{departamento}'" for departamento in DEPARTAMENTOS_TURISMO)

    # 2. Diccionario para almacenar los resultados
    dict_verif = {}

    # El siguiente proceso no es válido para la agrupación de COLOMBIA ya que esta siempre tiene datos: 
    if AGRUPACION in ['CONTINENTES', 'HUBS', 'TLCS', 'PAISES', 'DEPARTAMENTOS']:

        # 3. Verificación de exportaciones
        # Consultas para verificar datos de exportaciones en periodos cerrados y corridos, tanto totales como No Mineras (NME)
        
        # Totales
        # Cerrado
        query_verif_exportaciones_total_cerrado = f"""
        SELECT DISTINCT A.UNIDAD
        FROM DOCUMENTOS_COLOMBIA.EXPORTACIONES.ST_CATEGORIAS_CERRADO AS A
        WHERE A.AGRUPACION = '{AGRUPACION}' 
            AND A.UNIDAD IN ('{UNIDAD}')
            AND A.TABLA = 'TOTAL';
        """
        # Corrido
        query_verif_exportaciones_total_corrido = f"""
        SELECT DISTINCT A.UNIDAD
        FROM DOCUMENTOS_COLOMBIA.EXPORTACIONES.ST_CATEGORIAS_CORRIDO AS A
        WHERE A.AGRUPACION = '{AGRUPACION}' 
            AND A.UNIDAD IN ('{UNIDAD}')
            AND A.TABLA = 'TOTAL';
        """

        # NME
        # Cerrado
        query_verif_exportaciones_nme_cerrado = f"""
        SELECT DISTINCT A.UNIDAD
        FROM DOCUMENTOS_COLOMBIA.EXPORTACIONES.ST_CATEGORIAS_CERRADO AS A
        WHERE A.AGRUPACION = '{AGRUPACION}' 
            AND A.UNIDAD IN ('{UNIDAD}')
            AND A.TABLA = 'TIPOS'
            AND A.CATEGORIA = 'No Mineras';
        """
        # Corrido
        query_verif_exportaciones_nme_corrido = f"""
        SELECT DISTINCT A.UNIDAD
        FROM DOCUMENTOS_COLOMBIA.EXPORTACIONES.ST_CATEGORIAS_CORRIDO AS A
        WHERE A.AGRUPACION = '{AGRUPACION}' 
            AND A.UNIDAD IN ('{UNIDAD}')
            AND A.TABLA = 'TIPOS'
            AND A.CATEGORIA = 'No Mineras';
        """

        # Conteo de empresas
        # Cerrado
        query_verif_exportaciones_cuenta_cerrado = f"""
        SELECT A.NIT_EXPORTADOR, A.YEAR
            FROM DOCUMENTOS_COLOMBIA.EXPORTACIONES.ST_CONTEO_EMPRESAS_CERRADO AS A
            WHERE A.AGRUPACION = '{AGRUPACION}' 
                AND A.UNIDAD = '{UNIDAD}'
                AND A.VALOR_USD > {UMBRAL}
            ORDER BY A.YEAR ASC;
        """
        # Corrido
        query_verif_exportaciones_cuenta_corrido = f"""
        SELECT A.NIT_EXPORTADOR, A.YEAR
            FROM DOCUMENTOS_COLOMBIA.EXPORTACIONES.ST_CONTEO_EMPRESAS_CORRIDO AS A
            WHERE A.AGRUPACION = '{AGRUPACION}' 
                AND A.UNIDAD = '{UNIDAD}'
                AND A.VALOR_USD > {UMBRAL}
            ORDER BY A.YEAR ASC;
        """

        # Datos de empresas
        # Cerrado
        query_verif_exportaciones_empresas_cerrado = f"""
            SELECT A.CATEGORIA, 
                    A.RAZON_SOCIAL,
                    A.SECTOR_ESTRELLA,
                    A.SUMA_USD_T_1, 
                    A.SUMA_USD_T, 
                    A.DIFERENCIA_PORCENTUAL
            FROM DOCUMENTOS_COLOMBIA.EXPORTACIONES.ST_NIT_CERRADO AS A 
            WHERE A.AGRUPACION = '{AGRUPACION}' 
                AND A.UNIDAD = '{UNIDAD}'
            ORDER BY SUMA_USD_T DESC LIMIT 5;
        """
        # Corrido
        query_verif_exportaciones_empresas_corrido = f"""
            SELECT A.CATEGORIA, 
                    A.RAZON_SOCIAL,
                    A.SECTOR_ESTRELLA,
                    A.SUMA_USD_T_1, 
                    A.SUMA_USD_T, 
                    A.DIFERENCIA_PORCENTUAL
            FROM DOCUMENTOS_COLOMBIA.EXPORTACIONES.ST_NIT_CORRIDO AS A 
            WHERE A.AGRUPACION = '{AGRUPACION}' 
                AND A.UNIDAD = '{UNIDAD}'
            ORDER BY SUMA_USD_T DESC LIMIT 5;
        """

        # Ejecución de consultas y agregar al diccionario si los países o departamentos son válidos o no        
        df_export_total_cerrado = pd.DataFrame(session.sql(query_verif_exportaciones_total_cerrado).collect())
        df_export_total_corrido = pd.DataFrame(session.sql(query_verif_exportaciones_total_corrido).collect())
        df_nme_cerrado = pd.DataFrame(session.sql(query_verif_exportaciones_nme_cerrado).collect())
        df_nme_corrido = pd.DataFrame(session.sql(query_verif_exportaciones_nme_corrido).collect())
        df_conteo_cerrado = pd.DataFrame(session.sql(query_verif_exportaciones_cuenta_cerrado).collect())
        df_conteo_corrido = pd.DataFrame(session.sql(query_verif_exportaciones_cuenta_corrido).collect())
        df_empresas_cerrado = pd.DataFrame(session.sql(query_verif_exportaciones_empresas_cerrado).collect())
        df_empresas_corrido = pd.DataFrame(session.sql(query_verif_exportaciones_empresas_corrido).collect())


        # Verificar y agregar al diccionario dict_verif
        # Exportaciones Totales
        if df_export_total_cerrado.empty:
            dict_verif['exportaciones_totales_cerrado'] = "SIN DATOS DE EXPORTACIONES TOTALES CERRADO"
        else:
            dict_verif['exportaciones_totales_cerrado'] = "CON DATOS DE EXPORTACIONES TOTALES CERRADO"

        if df_export_total_corrido.empty:
            dict_verif['exportaciones_totales_corrido'] = "SIN DATOS DE EXPORTACIONES TOTALES CORRIDO"
        else:
            dict_verif['exportaciones_totales_corrido'] = "CON DATOS DE EXPORTACIONES TOTALES CORRIDO"

        # Exportaciones NME
        if df_nme_cerrado.empty:
            dict_verif['exportaciones_nme_cerrado'] = "SIN DATOS DE EXPORTACIONES NME CERRADO"
        else:
            dict_verif['exportaciones_nme_cerrado'] = "CON DATOS DE EXPORTACIONES NME CERRADO"

        if df_nme_corrido.empty:
            dict_verif['exportaciones_nme_corrido'] = "SIN DATOS DE EXPORTACIONES NME CORRIDO"
        else:
            dict_verif['exportaciones_nme_corrido'] = "CON DATOS DE EXPORTACIONES NME CORRIDO"

        # Conteo de empresas
        if df_conteo_cerrado.empty:
            dict_verif['exportaciones_conteo_cerrado'] = "SIN DATOS DE CONTEO CERRADO"
        else:
            dict_verif['exportaciones_conteo_cerrado'] = "CON DATOS DE CONTEO CERRADO"

        if df_conteo_corrido.empty:
            dict_verif['exportaciones_conteo_corrido'] = "SIN DATOS DE CONTEO CORRIDO"
        else:
            dict_verif['exportaciones_conteo_corrido'] = "CON DATOS DE CONTEO CORRIDO"

        # Empresas
        if df_empresas_cerrado.empty:
            dict_verif['exportaciones_empresas_cerrado'] = "SIN DATOS DE EMPRESAS CERRADO"
        else:
            dict_verif['exportaciones_empresas_cerrado'] = "CON DATOS DE EMPRESAS CERRADO"

        if df_empresas_corrido.empty:
            dict_verif['exportaciones_empresas_corrido'] = "SIN DATOS DE EMPRESAS CORRIDO"
        else:
            dict_verif['exportaciones_empresas_corrido'] = "CON DATOS DE EMPRESAS CORRIDO"


        # 4. Verificación de inversión 
        # Consultas para verificar datos de inversión extranjera directa (IED) e inversión colombiana en el exterior (ICE)
        if AGRUPACION in ['CONTINENTES', 'HUBS', 'TLCS', 'PAISES', 'COLOMBIA']:
            # IED
            # CERRADO
            query_verif_ied_cerrado = f"""
            SELECT DISTINCT A.UNIDAD
            FROM DOCUMENTOS_COLOMBIA.INVERSION.ST_PAISES_CERRADO AS A
            WHERE A.AGRUPACION = 'PAISES'
                AND A.UNIDAD NOT IN ('TOTAL')
                AND A.CATEGORIA = 'IED'
            """
            query_verif_ied_cerrado += f" AND A.UNIDAD IN ({PAISES_INVERSION_sql})"
            # CORRIDO
            query_verif_ied_corrido = f"""
            SELECT DISTINCT A.UNIDAD
            FROM DOCUMENTOS_COLOMBIA.INVERSION.ST_PAISES_CORRIDO AS A
            WHERE A.AGRUPACION = 'PAISES'
                AND A.UNIDAD NOT IN ('TOTAL')
                AND A.CATEGORIA = 'IED'
            """
            query_verif_ied_corrido += f" AND A.UNIDAD IN ({PAISES_INVERSION_sql})"

            # ICE
            # CERRADO
            query_verif_ice_cerrado = f"""
            SELECT DISTINCT A.UNIDAD
            FROM DOCUMENTOS_COLOMBIA.INVERSION.ST_PAISES_CERRADO AS A
            WHERE A.AGRUPACION = 'PAISES'
                AND A.UNIDAD NOT IN ('TOTAL')
                AND A.CATEGORIA = 'ICE'
            """
            query_verif_ice_cerrado += f" AND A.UNIDAD IN ({PAISES_INVERSION_sql})"
            # CORRIDO
            query_verif_ice_corrido = f"""
            SELECT DISTINCT A.UNIDAD
            FROM DOCUMENTOS_COLOMBIA.INVERSION.ST_PAISES_CORRIDO AS A
            WHERE A.AGRUPACION = 'PAISES'
                AND A.UNIDAD NOT IN ('TOTAL')
                AND A.CATEGORIA = 'ICE'
            """
            query_verif_ice_corrido += f" AND A.UNIDAD IN ({PAISES_INVERSION_sql})"

            # Ejecución de consultas y agregar al diccionario si los países son válidos o no

            # IED
            try:
                df_ied_cerrado = pd.DataFrame(session.sql(query_verif_ied_cerrado).collect())
                if df_ied_cerrado.empty:
                    dict_verif['ied_cerrado'] = "SIN DATOS DE IED CERRADO"
                else:
                    dict_verif['ied_cerrado'] = "CON DATOS DE IED CERRADO"
            except Exception as e:
                dict_verif['ied_cerrado'] = "SIN DATOS DE IED CERRADO"

            try:
                df_ied_corrido = pd.DataFrame(session.sql(query_verif_ied_corrido).collect())
                if df_ied_corrido.empty:
                    dict_verif['ied_corrido'] = "SIN DATOS DE IED CORRIDO"
                else:
                    dict_verif['ied_corrido'] = "CON DATOS DE IED CORRIDO"
            except Exception as e:
                dict_verif['ied_corrido'] = "SIN DATOS DE IED CORRIDO"

            # ICE
            try:
                df_ice_cerrado = pd.DataFrame(session.sql(query_verif_ice_cerrado).collect())
                if df_ice_cerrado.empty:
                    dict_verif['ice_cerrado'] = "SIN DATOS DE ICE CERRADO"
                else:
                    dict_verif['ice_cerrado'] = "CON DATOS DE ICE CERRADO"
            except Exception as e:
                dict_verif['ice_cerrado'] = "SIN DATOS DE ICE CERRADO"

            try:
                df_ice_corrido = pd.DataFrame(session.sql(query_verif_ice_corrido).collect())
                if df_ice_corrido.empty:
                    dict_verif['ice_corrido'] = "SIN DATOS DE ICE CORRIDO"
                else:
                    dict_verif['ice_corrido'] = "CON DATOS DE ICE CORRIDO"
            except Exception as e:
                dict_verif['ice_corrido'] = "SIN DATOS DE ICE CORRIDO"

        
        # 5. Verificación de turismo
        # Consultas para verificar datos de turismo en periodos cerrados y corridos, considerando agrupación por países o departamentos
        # Cerrado
        query_verif_turismo_cerrado = f"""
        SELECT A.PAIS_RESIDENCIA,
            FROM DOCUMENTOS_COLOMBIA.TURISMO.ST_PAISES_CERRADO AS A
        WHERE 1=1
        """
        if AGRUPACION in ['CONTINENTES', 'HUBS', 'TLCS', 'PAISES']:
            query_verif_turismo_cerrado += f" AND A.PAIS_RESIDENCIA IN ({PAISES_TURISMO_sql})"
        if AGRUPACION in ['DEPARTAMENTOS']:
            query_verif_turismo_cerrado += f" AND A.DPTO_HOSPEDAJE IN ({DEPARTAMENTOS_TURISMO_sql})"
        
        # Corrido
        query_verif_turismo_corrido = f"""
        SELECT A.PAIS_RESIDENCIA,
            FROM DOCUMENTOS_COLOMBIA.TURISMO.ST_PAISES_CORRIDO AS A
        WHERE 1=1
        """
        if AGRUPACION in ['CONTINENTES', 'HUBS', 'TLCS', 'PAISES']:
            query_verif_turismo_corrido += f" AND A.PAIS_RESIDENCIA IN ({PAISES_TURISMO_sql})"
        if AGRUPACION in ['DEPARTAMENTOS']:
            query_verif_turismo_corrido += f" AND A.DPTO_HOSPEDAJE IN ({DEPARTAMENTOS_TURISMO_sql})"
        
        # Ejecución de consultas y agregar al diccionario si los países son válidos o no
        # Turismo
        try:
            df_turismo_cerrado = pd.DataFrame(session.sql(query_verif_turismo_cerrado).collect())
            if df_turismo_cerrado.empty:
                dict_verif['turismo_cerrado'] = "SIN DATOS DE TURISMO CERRADO"
            else:
                dict_verif['turismo_cerrado'] = "CON DATOS DE TURISMO CERRADO"
        except Exception as e:
            dict_verif['turismo_cerrado'] = "SIN DATOS DE TURISMO CERRADO"

        try:
            df_turismo_corrido = pd.DataFrame(session.sql(query_verif_turismo_corrido).collect())
            if df_turismo_corrido.empty:
                dict_verif['turismo_corrido'] = "SIN DATOS DE TURISMO CORRIDO"
            else:
                dict_verif['turismo_corrido'] = "CON DATOS DE TURISMO CORRIDO"
        except Exception as e:
            dict_verif['turismo_corrido'] = "SIN DATOS DE TURISMO CORRIDO"
        
    ##############
    # CONECTIVIDAD
    ##############

    if AGRUPACION in ['DEPARTAMENTOS']:
        # Construir consulta
        query_conectividad = """
        SELECT A.AEROLINEA AS "Aerolínea",
            A.CIUDAD_ORIGEN AS "Ciudad Origen",
            A.CIUDAD_DESTINO AS "Ciudad Destino",
            A.FRECUENCIAS AS "Frecuencias",
            A.SEMANA AS "Semana de análisis"
        FROM DOCUMENTOS_COLOMBIA.TURISMO.CONECTIVIDAD AS A
        WHERE 1 = 1
        """
        # Agregar departamento
        query_conectividad += f" AND A.COD_DIVIPOLA_DEPARTAMENTO_DESTINO IN ({DEPARTAMENTOS_TURISMO_sql})"

        # Verificación
        try:
            # Ejecutar la consulta y recolectar los resultados en un DataFrame
            df_conectividad = pd.DataFrame(session.sql(query_conectividad).collect())
            
            # Verificar si el DataFrame está vacío
            if df_conectividad.empty:
                dict_verif['conectividad'] = "SIN DATOS DE CONECTIVIDAD"
            else:
                dict_verif['conectividad'] = "CON DATOS DE CONECTIVIDAD"
        except Exception as e:
            dict_verif['conectividad'] = "SIN DATOS DE CONECTIVIDAD"
    
    ###############
    # OPORTUNIDADES
    ###############
    if AGRUPACION in ['CONTINENTES', 'HUBS', 'TLCS', 'PAISES', 'DEPARTAMENTOS']:
        # Exportación
        query_oportunidades_exportacion = """
        SELECT DISTINCT A.CADENA,
            LOWER(A.SUBSECTOR) AS SUBSECTOR
        FROM DOCUMENTOS_COLOMBIA.EXPORTACIONES.OPORTUNIDADES AS A
        WHERE A.OPORTUNIDAD = 'Exportación' 
            AND A.CADENA NOT IN ('Turismo')
        """
        # Países y agrupaciones de países
        if AGRUPACION in ['CONTINENTES', 'HUBS', 'TLCS', 'PAISES']:
            query_oportunidades_exportacion += f" AND A.COD_PAIS IN ({PAISES_TURISMO_sql})"
        # Departamentos
        if AGRUPACION in ['DEPARTAMENTOS']:
            query_oportunidades_exportacion += f" AND A.COD_DIVIPOLA_DEPARTAMENTO IN ({DEPARTAMENTOS_TURISMO_sql})"
        # Order
        query_oportunidades_exportacion += f" ORDER BY 1, 2 ASC"

        # Verificación
        try:
            df_oportunidades_exportacion = pd.DataFrame(session.sql(query_oportunidades_exportacion).collect())
            if df_oportunidades_exportacion.empty:
                dict_verif['oportunidades_exportacion'] = "SIN OPORTUNIDADES"
            else:
                dict_verif['oportunidades_exportacion'] = "CON OPORTUNIDADES"
        except Exception as e:
            dict_verif['oportunidades_exportacion'] = "SIN OPORTUNIDADES"

        # Inversión
        query_oportunidades_ied = """
        SELECT DISTINCT A.CADENA,
            LOWER(A.SUBSECTOR) AS SUBSECTOR
        FROM DOCUMENTOS_COLOMBIA.EXPORTACIONES.OPORTUNIDADES AS A
        WHERE A.OPORTUNIDAD = 'IED'
        """
        # Países y agrupaciones de países
        if AGRUPACION in ['CONTINENTES', 'HUBS', 'TLCS', 'PAISES']:
            query_oportunidades_ied += f" AND A.COD_PAIS IN ({PAISES_TURISMO_sql})"
        # Departamentos
        if AGRUPACION in ['DEPARTAMENTOS']:
            query_oportunidades_ied += f" AND A.COD_DIVIPOLA_DEPARTAMENTO IN ({DEPARTAMENTOS_TURISMO_sql})"
        # Order
        query_oportunidades_ied += f" ORDER BY 1, 2 ASC"

        # Verificación
        try:
            df_oportunidades_inversion = pd.DataFrame(session.sql(query_oportunidades_ied).collect())
            if df_oportunidades_inversion.empty:
                dict_verif['oportunidades_inversion'] = "SIN OPORTUNIDADES"
            else:
                dict_verif['oportunidades_inversion'] = "CON OPORTUNIDADES"
        except Exception as e:
            dict_verif['oportunidades_inversion'] = "SIN OPORTUNIDADES"
        
        # Turismo
        query_oportunidades_turismo = """
        SELECT DISTINCT A.SECTOR,
            LOWER(A.SUBSECTOR) AS SUBSECTOR
        FROM DOCUMENTOS_COLOMBIA.EXPORTACIONES.OPORTUNIDADES AS A
        WHERE A.CADENA IN ('Turismo')
        """
        # Países y agrupaciones de países
        if AGRUPACION in ['CONTINENTES', 'HUBS', 'TLCS', 'PAISES']:
            query_oportunidades_turismo += f" AND A.COD_PAIS IN ({PAISES_TURISMO_sql})"
        # Departamentos
        if AGRUPACION in ['DEPARTAMENTOS']:
            query_oportunidades_turismo += f" AND A.COD_DIVIPOLA_DEPARTAMENTO IN ({DEPARTAMENTOS_TURISMO_sql})"
        # Order
        query_oportunidades_turismo += f" ORDER BY 1, 2 ASC"
        # Verificación
        try:
            df_oportunidades_turismo = pd.DataFrame(session.sql(query_oportunidades_turismo).collect())
            if df_oportunidades_turismo.empty:
                dict_verif['oportunidades_turismo'] = "SIN OPORTUNIDADES"
            else:
                dict_verif['oportunidades_turismo'] = "CON OPORTUNIDADES"
        except Exception as e:
            dict_verif['oportunidades_turismo'] = "SIN OPORTUNIDADES"

    
    #################
    # PESOS POR MEDIO
    #################
        # Mineros
        # Cerrado
        query_pesos_mineros_cerrado = f"""
        SELECT A.CATEGORIA, 
            A.SUMA_PESO_T_1, 
            A.SUMA_PESO_T, 
            A.DIFERENCIA_PORCENTUAL 
        FROM DOCUMENTOS_COLOMBIA.EXPORTACIONES.ST_CATEGORIAS_PESO_CERRADO AS A 
        WHERE A.TABLA = 'MEDIO MINERAS'
            AND A.AGRUPACION = '{AGRUPACION}' 
            AND A.UNIDAD IN ('{UNIDAD}');
        """
        # Verificación
        try:
            df_pesos_minero_cerrado = pd.DataFrame(session.sql(query_pesos_mineros_cerrado).collect())
            if df_pesos_minero_cerrado.empty:
                dict_verif['pesos_minero_cerrado'] = "SIN DATOS CERRADO"
            else:
                dict_verif['pesos_minero_cerrado'] = "CON DATOS CERRADO"
        except Exception as e:
            dict_verif['pesos_minero_cerrado'] = "SIN DATOS CERRADO"

        # Corrido
        query_pesos_mineros_corrido = f"""
        SELECT A.CATEGORIA, 
            A.SUMA_PESO_T_1, 
            A.SUMA_PESO_T, 
            A.DIFERENCIA_PORCENTUAL 
        FROM DOCUMENTOS_COLOMBIA.EXPORTACIONES.ST_CATEGORIAS_PESO_CORRIDO AS A 
        WHERE A.TABLA = 'MEDIO MINERAS'
            AND A.AGRUPACION = '{AGRUPACION}' 
            AND A.UNIDAD IN ('{UNIDAD}');
        """
        # Verificación
        try:
            df_pesos_minero_corrido = pd.DataFrame(session.sql(query_pesos_mineros_corrido).collect())
            if df_pesos_minero_corrido.empty:
                dict_verif['pesos_minero_corrido'] = "SIN DATOS CORRIDO"
            else:
                dict_verif['pesos_minero_corrido'] = "CON DATOS CORRIDO"
        except Exception as e:
            dict_verif['pesos_minero_corrido'] = "SIN DATOS CORRIDO"

        # No mineros
        # Cerrado
        query_pesos_no_mineros_cerrado = f"""
        SELECT A.CATEGORIA, 
            A.SUMA_PESO_T_1, 
            A.SUMA_PESO_T, 
            A.DIFERENCIA_PORCENTUAL 
        FROM DOCUMENTOS_COLOMBIA.EXPORTACIONES.ST_CATEGORIAS_PESO_CERRADO AS A 
        WHERE A.TABLA = 'MEDIO NO MINERAS'
            AND A.AGRUPACION = '{AGRUPACION}' 
            AND A.UNIDAD IN ('{UNIDAD}');
        """
        # Verificación
        try:
            df_pesos_no_minero_cerrado = pd.DataFrame(session.sql(query_pesos_no_mineros_cerrado).collect())
            if df_pesos_no_minero_cerrado.empty:
                dict_verif['pesos_no_minero_cerrado'] = "SIN DATOS CERRADO"
            else:
                dict_verif['pesos_no_minero_cerrado'] = "CON DATOS CERRADO"
        except Exception as e:
            dict_verif['pesos_no_minero_cerrado'] = "SIN DATOS CERRADO"

        # Corrido
        query_pesos_no_mineros_corrido = f"""
        SELECT A.CATEGORIA, 
            A.SUMA_PESO_T_1, 
            A.SUMA_PESO_T, 
            A.DIFERENCIA_PORCENTUAL 
        FROM DOCUMENTOS_COLOMBIA.EXPORTACIONES.ST_CATEGORIAS_PESO_CORRIDO AS A 
        WHERE A.TABLA = 'MEDIO NO MINERAS'
            AND A.AGRUPACION = '{AGRUPACION}' 
            AND A.UNIDAD IN ('{UNIDAD}');
        """
        # Verificación
        try:
            df_pesos_no_minero_corrido = pd.DataFrame(session.sql(query_pesos_no_mineros_corrido).collect())
            if df_pesos_no_minero_corrido.empty:
                dict_verif['pesos_no_minero_corrido'] = "SIN DATOS CORRIDO"
            else:
                dict_verif['pesos_no_minero_corrido'] = "CON DATOS CORRIDO"
        except Exception as e:
            dict_verif['pesos_no_minero_corrido'] = "SIN DATOS CORRIDO"
   
    # Colombia es válidos para los tres ejes siempre:
    if AGRUPACION == 'COLOMBIA':
        dict_verif['exportaciones_totales_cerrado'] = "CON DATOS DE EXPORTACIONES TOTALES CERRADO"
        dict_verif['exportaciones_totales_corrido'] = "CON DATOS DE EXPORTACIONES TOTALES CORRIDO"
        dict_verif['exportaciones_nme_cerrado'] = "CON DATOS DE EXPORTACIONES NME CERRADO"
        dict_verif['exportaciones_nme_corrido'] = "CON DATOS DE EXPORTACIONES NME CORRIDO"
        dict_verif['ied_cerrado'] = "CON DATOS DE IED CERRADO"
        dict_verif['ied_corrido'] = "CON DATOS DE IED CORRIDO"
        dict_verif['ice_cerrado'] = "CON DATOS DE ICE CERRADO"
        dict_verif['ice_corrido'] = "CON DATOS DE ICE CORRIDO"
        dict_verif['turismo_cerrado'] = "CON DATOS DE TURISMO CERRADO"
        dict_verif['turismo_corrido'] = "CON DATOS DE TURISMO CORRIDO"
        dict_verif['exportaciones_conteo_cerrado'] = "CON DATOS DE CONTEO CERRADO"
        dict_verif['exportaciones_conteo_corrido'] = "CON DATOS DE CONTEO CORRIDO"
        dict_verif['exportaciones_empresas_cerrado'] = "CON DATOS DE EMPRESAS CERRADO"
        dict_verif['exportaciones_empresas_corrido'] = "CON DATOS DE EMPRESAS CORRIDO"
        dict_verif['oportunidades_exportacion'] = "CON OPORTUNIDADES"
        dict_verif['oportunidades_inversion'] = "CON OPORTUNIDADES"
        dict_verif['oportunidades_turismo'] = "CON OPORTUNIDADES"
        dict_verif['conectividad'] = "SIN DATOS DE CONECTIVIDAD"
        dict_verif['pesos_minero_cerrado'] = "CON DATOS CERRADO"
        dict_verif['pesos_minero_corrido'] = "CON DATOS CORRIDO"
        dict_verif['pesos_no_minero_cerrado'] = "CON DATOS CERRADO"
        dict_verif['pesos_no_minero_corrido'] = "CON DATOS CORRIDO"

    return dict_verif


# Función auxiliar para personalizar estilos
def customize_style(style, font_name, font_size, font_color, bold=False):
    style.font.name = font_name
    style.font.size = font_size
    style.font.color.rgb = font_color
    style.font.bold = bold
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def estilos(doc: Document):
    """
    Define y aplica estilos personalizados al documento proporcionado.
    
    Args:
    doc (Document): El documento al que se añadirán los estilos.
    """
  
    # Personalizar el estilo 'Title'
    title_style = doc.styles['Title']
    customize_style(title_style, 'Century Gothic', Pt(16), RGBColor(0, 32, 96), bold=True)

    # Personalizar el estilo 'Heading 1'
    heading1_style = doc.styles['Heading 1']
    customize_style(heading1_style, 'Century Gothic', Pt(14), RGBColor(0, 32, 96), bold=True)

    # Personalizar el estilo 'Heading 2'
    heading2_style = doc.styles['Heading 2']
    customize_style(heading2_style, 'Century Gothic', Pt(12), RGBColor(0, 32, 96), bold=True)

    # Personalizar el estilo 'Heading 3'
    heading3_style = doc.styles['Heading 3']
    customize_style(heading3_style, 'Century Gothic', Pt(12), RGBColor(0, 32, 96), bold=True)

    # Personalizar el estilo 'Normal'
    normal_style = doc.styles['Normal']
    customize_style(normal_style, 'Century Gothic', Pt(11), RGBColor(0, 0, 0))

    # Personalizar el estilo 'Table'
    table_style = doc.styles['Table Grid']
    table_font = table_style.font
    table_font.name = 'Century Gothic'

    # Ajustar las márgenes del documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.44)  # Margen superior de 1.44 cm
        section.bottom_margin = Cm(2.5)  # Margen inferior de 2.5 cm
        section.left_margin = Cm(2.00)  # Margen izquierdo de 2.00 cm
        section.right_margin = Cm(2.00)  # Margen derecho de 2.00 cm


def add_heading(doc: Document, text: str, level: int, style: str):
    """
    Agrega un encabezado al documento.

    Args:
    doc (Document): El documento al que se añadirá el encabezado.
    text (str): El texto del encabezado.
    level (int): El nivel del encabezado.
    style (str): El estilo del encabezado.
    """
    doc.add_heading(text, level=level).style = doc.styles[style]


def add_paragraph(doc: Document, text: str, style: str):
    """
    Agrega un párrafo al documento.

    Args:
    doc (Document): El documento al que se añadirá el párrafo.
    text (str): El texto del párrafo.
    style (str): El estilo del párrafo.
    """
    p = doc.add_paragraph(text)
    p.style = doc.styles[style]


def set_cell_border(cell, **kwargs):
    """
    Establece los bordes de una celda.
    
    Uso:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "#00FF00", "space": "0"},
        left={"sz": 24, "val": "single", "color": "#0000FF", "space": "0"},
        right={"sz": 24, "val": "single", "color": "#000000", "space": "0"},
    )

    Args:
    cell: La celda a la que se aplicarán los bordes.
    kwargs: Un diccionario con las especificaciones para los bordes (tamaño, valor, color y espacio).
    """
    tcPr = cell._element.get_or_add_tcPr()

    # Verificar si ya existen bordes en la celda
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    # Aplicar los bordes según las especificaciones proporcionadas
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)

            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def add_table(doc: Document, dataframe: pd.DataFrame, style: str, font_size: int, fuente: str):
    """
    Agrega una tabla al documento a partir de un DataFrame y asegura que no se divida entre páginas.
    También agrega una nota al final con la fuente de los datos, asegurando que esté en la misma página que la tabla.

    Args:
    doc (Document): El documento al que se añadirá la tabla.
    dataframe (DataFrame): El DataFrame que se convertirá en tabla.
    style (str): El estilo de la tabla.
    font_size (int): El tamaño de la letra para los títulos y el contenido de la tabla.
    fuente (str): La fuente de los datos.
    """
    if not isinstance(dataframe, pd.DataFrame) or dataframe.empty:
        print(f"El valor proporcionado no es un DataFrame válido o está vacío: {dataframe}")
        return
    
    # Añadir la tabla al documento
    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    table.style = doc.styles[style]
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True  # Desactivar ajuste automático para controlar el ancho de las celdas

    # Configurar encabezados de la tabla
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(dataframe.columns):
        hdr_cells[i].text = str(column)
        paragraph = hdr_cells[i].paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_format = paragraph.paragraph_format
        paragraph_format.keep_with_next = True  # Mantener la cabecera con la siguiente fila
        paragraph_format.keep_together = True  # Mantener la cabecera en una sola página
        hdr_cells[i].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(font_size)
        set_cell_border(
            hdr_cells[i],
            top={"sz": 1, "val": "single", "color": "000000"},
            bottom={"sz": 1, "val": "single", "color": "000000"},
            left={"sz": 1, "val": "single", "color": "000000"},
            right={"sz": 1, "val": "single", "color": "000000"},
        )

    # Añadir filas de datos a la tabla
    for index, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)
            paragraph = row_cells[i].paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph_format = paragraph.paragraph_format
            paragraph_format.keep_with_next = True  # Mantener cada celda con la siguiente en la misma fila
            paragraph_format.keep_together = True  # Mantener el contenido de la celda en una sola página
            row_cells[i].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(font_size)
            set_cell_border(
                row_cells[i],
                top={"sz": 1, "val": "single", "color": "000000"},
                bottom={"sz": 1, "val": "single", "color": "000000"},
                left={"sz": 1, "val": "single", "color": "000000"},
                right={"sz": 1, "val": "single", "color": "000000"},
            )

    # Aplicar formato especial a la primera fila (cabecera)
    for cell in hdr_cells:
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), "#215E99")
        cell._element.get_or_add_tcPr().append(shading_elm)

    # Aplicar formato especial a la última fila y asegurar que se mantenga unida con la fuente
    last_row_cells = table.rows[-1].cells
    for cell in last_row_cells:
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.keep_with_next = True  # Mantener última fila unida con la fuente
        cell.paragraphs[0].runs[0].bold = True
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), "#DAE9F7")
        cell._element.get_or_add_tcPr().append(shading_elm)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)

    # Ajustar el ancho de las celdas para que la tabla ocupe todo el ancho de las márgenes del documento
    for row in table.rows:
        for cell in row.cells:
            cell.width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin

    # Agregar la fuente de los datos en un nuevo párrafo y asegurar que esté en la misma página que la tabla
    fuente_paragraph = doc.add_paragraph(f"Fuente: {fuente}", style='Normal')
    fuente_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    fuente_paragraph_format = fuente_paragraph.paragraph_format
    fuente_paragraph_format.keep_together = True  # Mantener la fuente en una sola página
    fuente_paragraph_format.keep_with_next = False  # No es necesario mantener con otro contenido después
    fuente_paragraph.space_before = Pt(12)  # Espacio antes del párrafo
    fuente_paragraph.space_after = Pt(12)  # Espacio después del párrafo
    fuente_paragraph.paragraph_format.left_indent = Cm(0.75)  # Indentación izquierda
    for run in fuente_paragraph.runs:
        run.font.size = Pt(9)  # Ajustar el tamaño de la fuente a 9 puntos
    
def add_table_resumen(doc: Document, dataframe: pd.DataFrame, style: str, font_size: int, fuente: str):
    """
    Agrega una tabla al documento a partir de un DataFrame y asegura que no se divida entre páginas.
    También agrega una nota al final con la fuente de los datos, asegurando que esté en la misma página que la tabla.

    Args:
    doc (Document): El documento al que se añadirá la tabla.
    dataframe (DataFrame): El DataFrame que se convertirá en tabla.
    style (str): El estilo de la tabla.
    font_size (int): El tamaño de la letra para los títulos y el contenido de la tabla.
    fuente (str): La fuente de los datos.
    """
    if not isinstance(dataframe, pd.DataFrame) or dataframe.empty:
        print(f"El valor proporcionado no es un DataFrame válido o está vacío: {dataframe}")
        return
    
    # Añadir la tabla al documento
    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    table.style = doc.styles[style]
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True  # Desactivar ajuste automático para controlar el ancho de las celdas
    
    # Configurar encabezados de la tabla
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(dataframe.columns):
        hdr_cells[i].text = str(column)
        paragraph = hdr_cells[i].paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_format = paragraph.paragraph_format
        paragraph_format.keep_with_next = True  # Mantener la cabecera con la siguiente fila
        paragraph_format.keep_together = True  # Mantener la cabecera en una sola página
        hdr_cells[i].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(font_size)
        set_cell_border(
            hdr_cells[i],
            top={"sz": 1, "val": "single", "color": "000000"},
            bottom={"sz": 1, "val": "single", "color": "000000"},
            left={"sz": 1, "val": "single", "color": "000000"},
            right={"sz": 1, "val": "single", "color": "000000"},
        )

    # Añadir filas de datos a la tabla
    for index, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)
            paragraph = row_cells[i].paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph_format = paragraph.paragraph_format
            paragraph_format.keep_with_next = True  # Mantener cada celda con la siguiente en la misma fila
            paragraph_format.keep_together = True  # Mantener el contenido de la celda en una sola página
            row_cells[i].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[i].paragraphs[0].runs[0].font.size = Pt(font_size)
            set_cell_border(
                row_cells[i],
                top={"sz": 1, "val": "single", "color": "000000"},
                bottom={"sz": 1, "val": "single", "color": "000000"},
                left={"sz": 1, "val": "single", "color": "000000"},
                right={"sz": 1, "val": "single", "color": "000000"},
            )

    # Aplicar formato especial a la primera fila (cabecera)
    for cell in hdr_cells:
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), "#215E99")
        cell._element.get_or_add_tcPr().append(shading_elm)

    # Ajustar el ancho de las celdas para que la tabla ocupe todo el ancho de las márgenes del documento
    for row in table.rows:
        for cell in row.cells:
            cell.width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin

    # Agregar la fuente de los datos en un nuevo párrafo y asegurar que esté en la misma página que la tabla
    fuente_paragraph = doc.add_paragraph(f"Fuente: {fuente}", style='Normal')
    fuente_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    fuente_paragraph_format = fuente_paragraph.paragraph_format
    fuente_paragraph_format.keep_together = True  # Mantener la fuente en una sola página
    fuente_paragraph_format.keep_with_next = False  # No es necesario mantener con otro contenido después
    fuente_paragraph.space_before = Pt(12)  # Espacio antes del párrafo
    fuente_paragraph.space_after = Pt(12)  # Espacio después del párrafo
    fuente_paragraph.paragraph_format.left_indent = Cm(0.75)  # Indentación izquierda
    for run in fuente_paragraph.runs:
        run.font.size = Pt(9)  # Ajustar el tamaño de la fuente a 9 puntos



def agregar_tabla_contenidos(new_doc, font_size=8):
    """
    Agrega una tabla de contenidos al documento proporcionado en el estilo 'Normal'.
    
    Args:
    new_doc (Document): El documento al que se añadirá la tabla de contenidos.
    font_size (int): Tamaño de la fuente para todo el documento en puntos (por defecto 12).
    """
    # Configurar el estilo de fuente y tamaño para todo el documento
    style = new_doc.styles['Normal']
    font = style.font
    font.size = Pt(font_size)

    # Crear un párrafo para el título de la tabla de contenidos y centrarlo
    para = new_doc.add_paragraph("Tabla de Contenidos")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.name = 'Century Gothic'
        run.font.size = Pt(11)
        run.bold = True
        run.underline = True
        run.font.color.rgb = RGBColor(0, 32, 96)

    # Crear un párrafo vacío para insertar el campo de la tabla de contenidos
    paragraph = new_doc.add_paragraph()
    paragraph.paragraph_format.space_before = Inches(0)
    paragraph.paragraph_format.space_after = Inches(0)
    run = paragraph.add_run()
    
    # Insertar el campo de la tabla de contenidos
    fldChar = OxmlElement('w:fldChar')  # crea un nuevo elemento
    fldChar.set(qn('w:fldCharType'), 'begin')  # establece el atributo en el elemento
 
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # establece el atributo en el elemento
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # cambia 1-3 dependiendo de los niveles de encabezado que necesites
 
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
 
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Haga clic derecho para actualizar el campo."
    fldChar3 = OxmlElement('w:updateFields')
    fldChar3.set(qn('w:val'), 'true')
    fldChar2.append(fldChar3)
 
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
 
    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)
 
    p_element = paragraph._p


def add_header_footer(doc: Document, header_image_left: str, footer_image: str, footer_text: str):
    """
    Agrega un encabezado con dos imágenes (una a la izquierda y otra a la derecha) y un pie de página con una imagen a la derecha, texto a la izquierda y el número de página en el centro.
    
    Args:
    doc (Document): El documento al que se añadirán el encabezado y el pie de página.
    footer_image (str): Ruta de la imagen para el pie de página.
    footer_text (str): Texto para el pie de página.
    """
    section = doc.sections[0]

    # Encabezado
    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=doc.sections[0].page_width)
    header_table.autofit = True

    # Imagen izquierda en el encabezado
    header_cell_left = header_table.cell(0, 0)
    header_paragraph_left = header_cell_left.paragraphs[0]
    header_run_left = header_paragraph_left.add_run()
    header_run_left.add_picture(header_image_left, width=Inches(2.5))

    # Pie de página
    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=3, width=doc.sections[0].page_width)
    footer_table.autofit = True

    # Texto izquierda en el pie de página
    footer_cell_left = footer_table.cell(0, 0)
    footer_paragraph_left = footer_cell_left.paragraphs[0]
    footer_paragraph_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    footer_run_left = footer_paragraph_left.add_run(footer_text)
    footer_run_left.font.size = Pt(8)

    # Número de página en el centro del pie de página
    footer_cell_center = footer_table.cell(0, 1)
    footer_paragraph_center = footer_cell_center.paragraphs[0]
    footer_paragraph_center.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run_center = footer_paragraph_center.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    footer_run_center._r.append(fldChar1)
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    footer_run_center._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    footer_run_center._r.append(fldChar2)
    # Establecer el tamaño de la fuente del número de página
    footer_run_center.font.size = Pt(8)

    # Imagen derecha en el pie de página
    footer_cell_right = footer_table.cell(0, 2)
    footer_paragraph_right = footer_cell_right.paragraphs[0]
    footer_paragraph_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    footer_run_right = footer_paragraph_right.add_run()
    footer_run_right.add_picture(footer_image, width=Inches(2.0))

def add_bullet_points(doc, bullet_points):
    """
    Agrega una lista de puntos de bala a un documento de Word con el texto justificado.

    Parámetros:
    - doc: El objeto Document de Word.
    - bullet_points: Lista de textos que irán en cada punto de bala.
    """
    for point in bullet_points:
        paragraph = doc.add_paragraph(style='List Bullet')
        run = paragraph.add_run(point)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Añadir configuración para justificar el texto si es necesario
        p = paragraph._element
        pPr = p.get_or_add_pPr()
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)


def agregar_oportunidades_al_documento(doc, data_dict, key):
    """
    Agrega las oportunidades de exportaciones o inversión a un documento Word.

    Esta función toma un documento Word y un diccionario de oportunidades de exportaciones,
    y agrega cada entrada al documento con la clave en negrilla seguida de los valores correspondientes.

    Parámetros:
    - doc (Document): El objeto Document de python-docx en el que se agregarán las oportunidades.
    - data_dict (dict): Un diccionario donde las claves son categorías de oportunidades y los valores
                        son las descripciones de las oportunidades.
    - key (str): La clave del diccionario que contiene las oportunidades de exportaciones o inversión.

    """
    # Acceder a las oportunidades en el diccionario
    oportunidades = data_dict.get(key, {})
    
    # Iterar sobre cada categoría y descripción en el diccionario de oportunidades
    for categoria, descripcion in oportunidades.items():
        # Agregar la categoría en negrilla
        p = doc.add_paragraph()
        run = p.add_run(f'{categoria}: ')
        run.bold = True
        
        # Agregar la descripción
        p.add_run(descripcion)

        # Justificar el contenido
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def obtener_parametros_documento(session):
    """
    Esta función extrae los parámetros específicos desde una base de datos en Snowflake
    y devuelve los resultados en un diccionario.

    Parámetros:
    - session: sesión de Snowflake.

    Retorna:
    - dict: Un diccionario con los parámetros solicitados.
    """
    # Definir la consulta SQL
    query = """
    SELECT A.PARAMETRO, A.VALOR
    FROM DOCUMENTOS_COLOMBIA.PARAMETROS.PARAMETROS AS A
    WHERE A.PARAMETRO IN ('Fecha de actualización', 'Año cerrado (T)', 'Año corrido texto (T)') 
        AND A.EJE IN ('Transversal', 'Exportaciones');
    """

    # Ejecutar la consulta y almacenar los resultados en un DataFrame de pandas
    data = pd.DataFrame(session.sql(query).collect())

    # Convertir el DataFrame en un diccionario
    parametros_dict = pd.Series(data.VALOR.values, index=data.PARAMETRO).to_dict()

    return parametros_dict

def create_document_continentes(tablas, file_path, titulo, header_image_left, footer_image, session, geo_params):
  
    doc = Document()
    estilos(doc)

    # Agregar encabezado y pie de página
    footer_text = """Calle 28 # 13ª - 15, Edificio CCI Pisos 35 - 36 | Bogotá, Colombia T: +57 (1) 560 0100 | info@procolombia.co | www.procolombia.co"""

    add_header_footer(doc, header_image_left, footer_image, footer_text)
        
    # Agregar el título principal del informe
    title_paragraph = doc.add_paragraph(f'TRES EJES CONTINENTES: {str(titulo).upper()}', style='Title')
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Agregar la fecha
    doc_params = obtener_parametros_documento(session)
    fecha = doc_params['Fecha de actualización']
    date_paragraph = doc.add_paragraph(f'ÚLTIMA ACTUALIZACIÓN: {fecha.upper()}', style='Title')
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
     
    # Agregar la tabla de contenido
    agregar_tabla_contenidos(doc, 10)
    doc.add_page_break()

    # Fuentes del documento:
    fuente_exportaciones = 'DANE-DIAN. Cálculos: ProColombia.'
    fuente_inversion = 'Banco de la República. Cálculos: ProColombia.'
    # Fuente condicional de Turismo con Venezuela
    if '850' in geo_params['PAISES_TURISMO_COD']:
        fuente_turismo = 'Migración Colombia: ProColombia. Nota: Los datos excluyen el registro de residentes venezolanos reportado por Migración Colombia (sin incluir la estimación del MinCIT). Tampoco se incluyen las cifras de colombianos residentes en el exterior ni de cruceristas.'
    else:
        fuente_turismo = 'Migración Colombia: ProColombia. Nota: Los datos no incluyen las cifras de colombianos residentes en el exterior ni de cruceristas.'

    # Disclaimer:
    disclaimer = "La información contenida en este documento es de orientación y guía general. En ningún caso, ProColombia, ni sus empleados, son responsables ante usted o cualquier otra persona por las decisiones o acciones que pueda tomar en relación con la información proporcionada, por lo cual debe tomarse como de carácter referencial únicamente."

    # Obtener parámetos de veríficación
    dict_verificacion = verif_ejes(session, geo_params)

    #########
    # RESUMEN
    #########
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO') or (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO') or (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO') or (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
        add_heading(doc, 'Resumen', level=2, style='Heading 1')
        # Exportaciones
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_heading(doc, 'Exportaciones', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_expo"]), 'Table Grid', 10, fuente_exportaciones)

            # Lista para los bullets
            bullets_exportaciones = []

            # Agregar los bullets según las condiciones
            if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
                if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b1_corrido"])
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b2_corrido"])
                else:
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_cerrado"])
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_cerrado"])
                # Siempre se incluye el bullet 3 de empresas:
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b3_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b3_corrido"])
            elif (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_corrido"])
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_corrido"])
                # Siempre se incluye el bullet 3 de empresas:
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b3_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b3_corrido"])

            # Usar la función para agregar los bullets al documento
            add_bullet_points(doc, bullets_exportaciones)


        # Inversión
        if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO') or (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
            add_heading(doc, 'Inversión', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_inv"]), 'Table Grid', 10, fuente_inversion)
            
            # Lista para los bullets de inversión
            bullets_inversion = []

            # Agregar los bullets de inversión según las condiciones
            if dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO':
                if dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO':
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b1_cerrado"] + " " + tablas["RESUMEN"]["texto_inversion_b1_corrido"])
                else:
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b1_cerrado"])
            elif dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO':
                bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b1_corrido"])

            if dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO':
                if dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO':
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b2_cerrado"] + " " + tablas["RESUMEN"]["texto_inversion_b2_corrido"])
                else:
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b2_cerrado"])
            elif dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO':
                bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b2_corrido"])

            # Usar la función para agregar los bullets de inversión al documento
            add_bullet_points(doc, bullets_inversion)


        # Turismo
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            doc.add_page_break()
            add_heading(doc, 'Turismo', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_tur"]), 'Table Grid', 10, fuente_turismo)
            
            # Agregar los bullets de turismo según las condiciones
            if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
                add_bullet_points(doc, [tablas["RESUMEN"]["texto_turismo_b1_cerrado"]])    
            if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
                add_bullet_points(doc, [tablas["RESUMEN"]["texto_turismo_b2_corrido"]])

        # Salto de página
        doc.add_page_break()

    ###############
    # Exportaciones
    ###############
    add_heading(doc, 'Exportaciones', level=2, style='Heading 1')
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
        
        #####################
        # Tipo de exportación
        #####################
        add_heading(doc, 'Tipo de exportación', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TIPOS"]["ST_CATEGORIAS_CERRADO"]), 'Table Grid', 10, fuente_exportaciones)
        # Año corrido
        if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TIPOS"]["ST_CATEGORIAS_CORRIDO"]), 'Table Grid', 10, fuente_exportaciones)

        #####################################
        # Exportaciones no minero-energéticas
        #####################################
        if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO') or (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
            add_heading(doc, 'Exportaciones no minero-energéticas', level=2, style='Heading 1')

            ###########
            #  Destinos
            ###########
            add_heading(doc, 'Destinos', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'): 
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["PAIS"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["PAIS"]), 'Table Grid', 10, fuente_exportaciones)
            
            ########################
            # Departamento de origen
            ########################
            add_heading(doc, 'Departamento de origen', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["DEPARTAMENTOS"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["DEPARTAMENTOS"]), 'Table Grid', 10, fuente_exportaciones)

            ########
            # Sector
            ########
            add_heading(doc, 'Sector', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["SECTORES"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["SECTORES"]), 'Table Grid', 10, fuente_exportaciones)

            #############
            # SubSectores
            #############
            add_heading(doc, 'Subsector', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["SUBSECTORES"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["SUBSECTORES"]), 'Table Grid', 10, fuente_exportaciones)

            ####################
            # Conteo de Empresas
            ####################
            if (dict_verificacion['exportaciones_conteo_cerrado'] == 'CON DATOS DE CONTEO CERRADO') or (dict_verificacion['exportaciones_conteo_corrido'] == 'CON DATOS DE CONTEO CORRIDO'):
                add_heading(doc, 'Empresas', level=3, style='Heading 2')
                # Año cerrado
                if (dict_verificacion['exportaciones_conteo_cerrado'] == 'CON DATOS DE CONTEO CERRADO'):
                    year_cerrado = doc_params['Año cerrado (T)']
                    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
                    conteo_cerrado = tablas['CONTEO EMPRESAS']['CERRADO'][0]
                    run_cerrado = paragraph_cerrado.add_run(f"{conteo_cerrado} empresas")
                    run_cerrado.bold = True
                elif (dict_verificacion['exportaciones_conteo_cerrado'] == 'SIN DATOS DE CONTEO CERRADO'):
                    year_cerrado = doc_params['Año cerrado (T)']
                    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
                    conteo_cerrado = "0"
                    run_cerrado = paragraph_cerrado.add_run(f"{conteo_cerrado} empresas")
                    run_cerrado.bold = True                    
                # Año corrido
                if (dict_verificacion['exportaciones_conteo_corrido'] == 'CON DATOS DE CONTEO CORRIDO'):
                    year_corrido = doc_params['Año corrido texto (T)']
                    paragraph_corrido = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
                    conteo_corrido = tablas['CONTEO EMPRESAS']['CORRIDO'][0]
                    run_corrido = paragraph_corrido.add_run(f"{conteo_corrido} empresas")
                    run_corrido.bold = True
                elif (dict_verificacion['exportaciones_conteo_corrido'] == 'SIN DATOS DE CONTEO CORRIDO'):
                    year_corrido = doc_params['Año corrido texto (T)']
                    paragraph_corrido = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
                    conteo_corrido = "0"
                    run_corrido = paragraph_corrido.add_run(f"{conteo_corrido} empresas")
                    run_corrido.bold = True

            ################
            # Datos empresas
            ################
            if (dict_verificacion['exportaciones_empresas_cerrado'] == 'CON DATOS DE EMPRESAS CERRADO') or (dict_verificacion['exportaciones_empresas_corrido'] == 'CON DATOS DE EMPRESAS CORRIDO'):
                add_heading(doc, 'Información de Empresas', level=3, style='Heading 2')
                # Año cerrado
                if (dict_verificacion['exportaciones_empresas_cerrado'] == 'CON DATOS DE EMPRESAS CERRADO'):
                    add_table(doc, pd.DataFrame(tablas['EMPRESAS']['ST_NIT_CERRADO']), 'Table Grid', 9, fuente_exportaciones)
                # Año corrido
                if (dict_verificacion['exportaciones_empresas_corrido'] == 'CON DATOS DE EMPRESAS CORRIDO'):
                    add_table(doc, pd.DataFrame(tablas['EMPRESAS']['ST_NIT_CORRIDO']), 'Table Grid', 9, fuente_exportaciones)
            
            ##############################
            # Oportunidades de exportación
            ##############################
            if (dict_verificacion['oportunidades_exportacion'] == "CON OPORTUNIDADES"):
                add_heading(doc, 'Oportunidades de exportación identificadas', level=3, style='Heading 2')
                agregar_oportunidades_al_documento(doc, tablas, 'OPORTUNIDADES_EXPORTACIONES')
            
        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('El continente no registra datos de exportaciones.')


    ###########
    # Inversión
    ###########
    add_heading(doc, 'Inversión', level=2, style='Heading 1')
    if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO') or (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
        
        #####
        # IED
        #####
        # Total
        if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO'):
            add_heading(doc, 'IED', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO'):
                add_table(doc, pd.DataFrame(tablas['IED TOTAL']['ied_cerrado_total']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['IED TOTAL']['ied_corrido_total']), 'Table Grid', 10, fuente_inversion)
            # Países
            add_heading(doc, 'IED - Países', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO'):
                add_table(doc, pd.DataFrame(tablas['IED PAISES']['ied_cerrado']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['IED PAISES']['ied_corrido']), 'Table Grid', 10, fuente_inversion)

        #####
        # ICE
        #####
        # Total
        if (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
            add_heading(doc, 'ICE', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO'):
                add_table(doc, pd.DataFrame(tablas['ICE TOTAL']['ice_cerrado_total']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['ICE TOTAL']['ice_corrido_total']), 'Table Grid', 10, fuente_inversion)
            # Países
            add_heading(doc, 'ICE - Países', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO'):
                add_table(doc, pd.DataFrame(tablas['ICE PAISES']['ice_cerrado']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['ICE PAISES']['ice_corrido']), 'Table Grid', 10, fuente_inversion)
        
        ############################
        # Oportunidades de inversión
        ############################
        if (dict_verificacion['oportunidades_inversion'] == "CON OPORTUNIDADES"):
            add_heading(doc, 'Oportunidades de inversión identificadas', level=3, style='Heading 2')
            agregar_oportunidades_al_documento(doc, tablas, 'OPORTUNIDADES_INVERSION')

        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('El continente no registra datos de inversión.')
    

    #########
    # Turismo
    #########
    add_heading(doc, 'Turismo', level=2, style='Heading 1')
    if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
        
        ########
        # Países
        ########
        add_heading(doc, 'Países', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["PAIS_RESIDENCIA"]), 'Table Grid', 10, fuente_turismo)
        # Año corrido
        if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CORRIDO"]["PAIS_RESIDENCIA"]), 'Table Grid', 10, fuente_turismo)

        ###############
        # Departamentos
        ###############
        add_heading(doc, 'Departamentos', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["DPTO_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)
        # Año corrido
        if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CORRIDO"]["DPTO_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)

        ############
        # Municipios
        ############
        add_heading(doc, 'Municipios', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["CIUDAD_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)
        # Año corrido
        if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CORRIDO"]["CIUDAD_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)

        ########
        # Género
        ########
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_heading(doc, 'Género', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["DESCRIPCION_GENERO"]), 'Table Grid', 10, fuente_turismo)

        #########
        # Motivos
        #########
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_heading(doc, 'Motivos', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["MOVC_NOMBRE"]), 'Table Grid', 10, fuente_turismo)

        ###############
        # OPORTUNIDADES
        ###############
        if (dict_verificacion['oportunidades_turismo'] == "CON OPORTUNIDADES"):
            add_heading(doc, 'Productos vacacionales que se promocionan', level=3, style='Heading 2')
            # Principales
            p1 = doc.add_paragraph()
            run1 = p1.add_run("Principales: ")
            run1.bold = True
            p1.add_run(tablas['TURISMO_PRINCIPAL'])
            p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Nichos
            p2 = doc.add_paragraph()
            run2 = p2.add_run("Nichos: ")
            run2.bold = True
            p2.add_run(tablas['TURISMO_NICHOS'])
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('El continente no registra datos de turismo.')

    ###########
    # Logística
    ###########
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
        add_heading(doc, 'Logística', level=2, style='Heading 1')
    #######
    # Pesos
    #######
        add_heading(doc, 'Pesos', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
            add_table(doc, pd.DataFrame(tablas['TIPOS PESO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
        # Año corrido
        if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_table(doc, pd.DataFrame(tablas['TIPOS PESO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)
        # Mineros
        if (dict_verificacion['pesos_minero_cerrado'] == 'CON DATOS CERRADO') or (dict_verificacion['pesos_minero_corrido'] == 'CON DATOS CORRIDO'):
            add_heading(doc, 'Pesos por medio de transporte: exportaciones mineras', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO MINERO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO MINERO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)
        # No Mineros
        if (dict_verificacion['pesos_no_minero_cerrado'] == 'CON DATOS CERRADO') or (dict_verificacion['pesos_no_minero_corrido'] == 'CON DATOS CORRIDO'):
            add_heading(doc, 'Pesos por medio de transporte: exportaciones no mineras', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO NO MINERO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO NO MINERO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)
        # Salto de página
        doc.add_page_break()
    
    #######
    # Anexo
    #######
    add_heading(doc, 'Anexo: Países considerados', level=2, style='Heading 1')
    # Agregar el párrafo introductorio
    paragraph_intro = doc.add_paragraph(
        f'El presente documento muestra los datos agregados para los Tres Ejes de negocio de ProColombia para el continente de {str(titulo).capitalize()}. Se incluye información de los siguientes países:')
    paragraph_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Agregar el párrafo con los países, añadir un punto al final y justificarlo
    paises_anexo_con_punto = geo_params['PAISES_ANEXO'] + '.'
    paragraph_paises = doc.add_paragraph(paises_anexo_con_punto)
    paragraph_paises.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ############
    # Disclaimer
    ############
    # Salto de página
    doc.add_page_break()
    # Agregar saltos de línea para centrar el texto verticalmente
    for _ in range(12): 
        doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
    # Agregar el texto del disclaimer
    paragraph_disclaimer = doc.add_paragraph(disclaimer)
    # Formato del texto
    run = paragraph_disclaimer.runs[0]
    run.font.name = 'Century Gothic'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 32, 96)
    run.bold = True
    paragraph_disclaimer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Guardar el documento
    doc.save(file_path)

def create_document_colombia(tablas, file_path, header_image_left, footer_image, session, geo_params):
  
    doc = Document()
    estilos(doc)

    # Agregar encabezado y pie de página
    footer_text = """Calle 28 # 13ª - 15, Edificio CCI Pisos 35 - 36 | Bogotá, Colombia T: +57 (1) 560 0100 | info@procolombia.co | www.procolombia.co"""

    add_header_footer(doc, header_image_left, footer_image, footer_text)
        
    # Agregar el título principal del informe
    title_paragraph = doc.add_paragraph(f'TRES EJES COLOMBIA', style='Title')
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Agregar la fecha
    doc_params = obtener_parametros_documento(session)
    fecha = doc_params['Fecha de actualización']
    date_paragraph = doc.add_paragraph(f'ÚLTIMA ACTUALIZACIÓN: {fecha.upper()}', style='Title')
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
     
    # Agregar la tabla de contenido
    agregar_tabla_contenidos(doc, 10)
    doc.add_page_break()

    # Fuentes del documento:
    fuente_exportaciones = 'DANE-DIAN. Cálculos: ProColombia.'
    fuente_inversion = 'Banco de la República. Cálculos: ProColombia.'
    fuente_turismo = 'Migración Colombia: ProColombia. Nota: Los datos excluyen el registro de residentes venezolanos reportado por Migración Colombia (sin incluir la estimación del MinCIT). Tampoco se incluyen las cifras de colombianos residentes en el exterior ni de cruceristas.'

    # Disclaimer:
    disclaimer = "La información contenida en este documento es de orientación y guía general. En ningún caso, ProColombia, ni sus empleados, son responsables ante usted o cualquier otra persona por las decisiones o acciones que pueda tomar en relación con la información proporcionada, por lo cual debe tomarse como de carácter referencial únicamente."

    # Obtener parámetos de veríficación
    dict_verificacion = verif_ejes(session, geo_params)

    #########
    # RESUMEN
    #########
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO') or (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO') or (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO') or (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
        add_heading(doc, 'Resumen', level=2, style='Heading 1')
        # Exportaciones
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_heading(doc, 'Exportaciones', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_expo"]), 'Table Grid', 10, fuente_exportaciones)

            # Lista para los bullets
            bullets_exportaciones = []

            # Agregar los bullets según las condiciones
            if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
                if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b1_corrido"])
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b2_corrido"])
                else:
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_cerrado"])
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_cerrado"])
                # Siempre se incluye el bullet 3 de empresas:
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b3_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b3_corrido"])
            elif (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_corrido"])
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_corrido"])
                # Siempre se incluye el bullet 3 de empresas:
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b3_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b3_corrido"])

            # Usar la función para agregar los bullets al documento
            add_bullet_points(doc, bullets_exportaciones)

        # Inversión
        if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO') or (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
            add_heading(doc, 'Inversión', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_inv"]), 'Table Grid', 10, fuente_inversion)
            
            # Lista para los bullets de inversión
            bullets_inversion = []

            # Agregar los bullets de inversión según las condiciones
            if dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO':
                if dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO':
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b1_cerrado"] + " " + tablas["RESUMEN"]["texto_inversion_b1_corrido"])
                else:
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b1_cerrado"])
            elif dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO':
                bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b1_corrido"])

            if dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO':
                if dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO':
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b2_cerrado"] + " " + tablas["RESUMEN"]["texto_inversion_b2_corrido"])
                else:
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b2_cerrado"])
            elif dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO':
                bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b2_corrido"])

            # Usar la función para agregar los bullets de inversión al documento
            add_bullet_points(doc, bullets_inversion)


        # Turismo
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_heading(doc, 'Turismo', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_tur"]), 'Table Grid', 10, fuente_turismo)

            # Agregar los bullets de turismo según las condiciones
            if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
                add_bullet_points(doc, [tablas["RESUMEN"]["texto_turismo_b1_cerrado"]])    
            if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
                add_bullet_points(doc, [tablas["RESUMEN"]["texto_turismo_b2_corrido"]])

        # Salto de página
        doc.add_page_break()

    ###############
    # Exportaciones
    ###############
    add_heading(doc, 'Exportaciones', level=2, style='Heading 1')
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
        
        #####################
        # Tipo de exportación
        #####################
        add_heading(doc, 'Tipo de exportación', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TIPOS"]["ST_CATEGORIAS_CERRADO"]), 'Table Grid', 10, fuente_exportaciones)
        # Año corrido
        if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TIPOS"]["ST_CATEGORIAS_CORRIDO"]), 'Table Grid', 10, fuente_exportaciones)

        #####################################
        # Exportaciones no minero-energéticas
        #####################################
        if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO') or (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
            add_heading(doc, 'Exportaciones no minero-energéticas', level=2, style='Heading 1')

            ###########
            #  Destinos
            ###########
            add_heading(doc, 'Destinos', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'): 
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["PAIS"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["PAIS"]), 'Table Grid', 10, fuente_exportaciones)
            
            ########################
            # Departamento de origen
            ########################
            add_heading(doc, 'Departamento de origen', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["DEPARTAMENTOS"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["DEPARTAMENTOS"]), 'Table Grid', 10, fuente_exportaciones)

            ########
            # Sector
            ########
            add_heading(doc, 'Sector', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["SECTORES"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["SECTORES"]), 'Table Grid', 10, fuente_exportaciones)

            #############
            # SubSectores
            #############
            add_heading(doc, 'Subsector', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["SUBSECTORES"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["SUBSECTORES"]), 'Table Grid', 10, fuente_exportaciones)

            ####################
            # Conteo de Empresas
            ####################
            if (dict_verificacion['exportaciones_conteo_cerrado'] == 'CON DATOS DE CONTEO CERRADO') or (dict_verificacion['exportaciones_conteo_corrido'] == 'CON DATOS DE CONTEO CORRIDO'):
                add_heading(doc, 'Empresas', level=3, style='Heading 2')
                # Año cerrado
                if (dict_verificacion['exportaciones_conteo_cerrado'] == 'CON DATOS DE CONTEO CERRADO'):
                    year_cerrado = doc_params['Año cerrado (T)']
                    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
                    conteo_cerrado = tablas['CONTEO EMPRESAS']['CERRADO'][0]
                    run_cerrado = paragraph_cerrado.add_run(f"{conteo_cerrado} empresas")
                    run_cerrado.bold = True
                elif (dict_verificacion['exportaciones_conteo_cerrado'] == 'SIN DATOS DE CONTEO CERRADO'):
                    year_cerrado = doc_params['Año cerrado (T)']
                    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
                    conteo_cerrado = "0"
                    run_cerrado = paragraph_cerrado.add_run(f"{conteo_cerrado} empresas")
                    run_cerrado.bold = True                    
                # Año corrido
                if (dict_verificacion['exportaciones_conteo_corrido'] == 'CON DATOS DE CONTEO CORRIDO'):
                    year_corrido = doc_params['Año corrido texto (T)']
                    paragraph_corrido = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
                    conteo_corrido = tablas['CONTEO EMPRESAS']['CORRIDO'][0]
                    run_corrido = paragraph_corrido.add_run(f"{conteo_corrido} empresas")
                    run_corrido.bold = True
                elif (dict_verificacion['exportaciones_conteo_corrido'] == 'SIN DATOS DE CONTEO CORRIDO'):
                    year_corrido = doc_params['Año corrido texto (T)']
                    paragraph_corrido = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
                    conteo_corrido = "0"
                    run_corrido = paragraph_corrido.add_run(f"{conteo_corrido} empresas")
                    run_corrido.bold = True

            ################
            # Datos empresas
            ################
            if (dict_verificacion['exportaciones_empresas_cerrado'] == 'CON DATOS DE EMPRESAS CERRADO') or (dict_verificacion['exportaciones_empresas_corrido'] == 'CON DATOS DE EMPRESAS CORRIDO'):
                add_heading(doc, 'Información de Empresas', level=3, style='Heading 2')
                # Año cerrado
                if (dict_verificacion['exportaciones_empresas_cerrado'] == 'CON DATOS DE EMPRESAS CERRADO'):
                    add_table(doc, pd.DataFrame(tablas['EMPRESAS']['ST_NIT_CERRADO']), 'Table Grid', 9, fuente_exportaciones)
                # Año corrido
                if (dict_verificacion['exportaciones_empresas_corrido'] == 'CON DATOS DE EMPRESAS CORRIDO'):
                    add_table(doc, pd.DataFrame(tablas['EMPRESAS']['ST_NIT_CORRIDO']), 'Table Grid', 9, fuente_exportaciones)
            
            ##############################
            # Oportunidades de exportación
            ##############################
            if (dict_verificacion['oportunidades_exportacion'] == "CON OPORTUNIDADES"):
                add_heading(doc, 'Oportunidades de exportación identificadas', level=3, style='Heading 2')
                agregar_oportunidades_al_documento(doc, tablas, 'OPORTUNIDADES_EXPORTACIONES')
        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('Colombia no registra datos de exportaciones.')
    

    ###########
    # Inversión
    ###########
    add_heading(doc, 'Inversión', level=2, style='Heading 1')
    if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO') or (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
        
        #####
        # IED
        #####
        # Total
        if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO'):
            add_heading(doc, 'IED', level=3, style='Heading 2')
            # Países
            add_heading(doc, 'IED - Países', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO'):
                add_table(doc, pd.DataFrame(tablas['IED PAISES']['ied_cerrado']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['IED PAISES']['ied_corrido']), 'Table Grid', 10, fuente_inversion)
            # Actividades
            add_heading(doc, 'IED - Actividades', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO'):
                add_table(doc, pd.DataFrame(tablas['IED ACTIVIDADES COLOMBIA']['ied_cerrado']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['IED ACTIVIDADES COLOMBIA']['ied_corrido']), 'Table Grid', 10, fuente_inversion)

        #####
        # ICE
        #####
        # Total
        if (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
            add_heading(doc, 'ICE', level=3, style='Heading 2')
            # Países
            add_heading(doc, 'ICE - Países', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO'):
                add_table(doc, pd.DataFrame(tablas['ICE PAISES']['ice_cerrado']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['ICE PAISES']['ice_corrido']), 'Table Grid', 10, fuente_inversion)
        
        ############################
        # Oportunidades de inversión
        ############################
        if (dict_verificacion['oportunidades_inversion'] == "CON OPORTUNIDADES"):
            add_heading(doc, 'Oportunidades de inversión identificadas', level=3, style='Heading 2')
            agregar_oportunidades_al_documento(doc, tablas, 'OPORTUNIDADES_INVERSION')
        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('Colombia no registra datos de inversión.')

    #########
    # Turismo
    #########
    add_heading(doc, 'Turismo', level=2, style='Heading 1')
    if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
        
        ########
        # Países
        ########
        add_heading(doc, 'Países', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["PAIS_RESIDENCIA"]), 'Table Grid', 10, fuente_turismo)
        # Año corrido
        if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CORRIDO"]["PAIS_RESIDENCIA"]), 'Table Grid', 10, fuente_turismo)

        ###############
        # Departamentos
        ###############
        add_heading(doc, 'Departamentos', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["DPTO_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)
        # Año corrido
        if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CORRIDO"]["DPTO_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)

        ############
        # Municipios
        ############
        add_heading(doc, 'Municipios', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["CIUDAD_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)
        # Año corrido
        if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CORRIDO"]["CIUDAD_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)

        ########
        # Género
        ########
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_heading(doc, 'Género', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["DESCRIPCION_GENERO"]), 'Table Grid', 10, fuente_turismo)

        #########
        # Motivos
        #########
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_heading(doc, 'Motivos', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["MOVC_NOMBRE"]), 'Table Grid', 10, fuente_turismo)

        ###############
        # OPORTUNIDADES
        ###############
        if (dict_verificacion['oportunidades_turismo'] == "CON OPORTUNIDADES"):
            add_heading(doc, 'Productos vacacionales que se promocionan', level=3, style='Heading 2')
            # Principales
            p1 = doc.add_paragraph()
            run1 = p1.add_run("Principales: ")
            run1.bold = True
            p1.add_run(tablas['TURISMO_PRINCIPAL'])
            p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Nichos
            p2 = doc.add_paragraph()
            run2 = p2.add_run("Nichos: ")
            run2.bold = True
            p2.add_run(tablas['TURISMO_NICHOS'])
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('Colombia no registra datos de turismo.')
    
    ###########
    # Logística
    ###########
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
        add_heading(doc, 'Logística', level=2, style='Heading 1')
    #######
    # Pesos
    #######
        add_heading(doc, 'Pesos', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
            add_table(doc, pd.DataFrame(tablas['TIPOS PESO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
        # Año corrido
        if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_table(doc, pd.DataFrame(tablas['TIPOS PESO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)
        # Mineros
        if (dict_verificacion['pesos_minero_cerrado'] == 'CON DATOS CERRADO') or (dict_verificacion['pesos_minero_corrido'] == 'CON DATOS CORRIDO'):
            add_heading(doc, 'Pesos por medio de transporte: exportaciones mineras', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO MINERO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO MINERO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)
        # No Mineros
        if (dict_verificacion['pesos_no_minero_cerrado'] == 'CON DATOS CERRADO') or (dict_verificacion['pesos_no_minero_corrido'] == 'CON DATOS CORRIDO'):
            add_heading(doc, 'Pesos por medio de transporte: exportaciones no mineras', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO NO MINERO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO NO MINERO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)
        # Usar la función para agregar los bullets al documento
        add_heading(doc, 'Principales puntos de llegada:', level=3, style='Heading 2')
        llegada1 = "Colombia cuenta con una ubicación privilegiada, situada en el punto focal de la actividad marítima por su cercanía al Canal de Panamá y en el cruce de las principales vías de comunicación del comercio mundial. Es punto de conexión estratégico entre Norte y Sur América, y entre la Costa Este de EE.UU y Asia, es una oportunidad como Plataforma de Intercambio Comercial Andino."
        llegada2 = "En los puertos y aeropuertos nacionales, se enlazan las más importantes navieras y aerolíneas del mundo, siendo punto estratégico en la conectividad global.  Colombia se conecta con más de 450 aeropuertos del mundo. El Aeropuerto Internacional el Dorado localizado en la ciudad de Bogotá, se ubica como el de mayor importancia en el país. En 2017 por los aeropuertos colombianos se movilizaron más de un millón de toneladas de carga internacional."
        llegada3 = "Además, Colombia se conecta con más de 600 puertos a nivel mundial. En el 2017, Colombia movilizó por sus zonas portuarias más de 205 millones de toneladas de carga."
        llegada4 = "Más de 4.314 rutas de exportación en servicio regular directas y con conexión prestadas por 32 empresas de transporte marítimo con destino a más de 661 puertos en el mundo. "
        llegada5 = "Más de 2.045 aéreas prestadas por 29 aerolíneas con cupos en aviones cargueros o aviones de pasajeros con cupo para carga con destino a más de 457 ciudad en el mundo."
        bullets_llegadas = [llegada1, llegada2, llegada3, llegada4, llegada5]
        add_bullet_points(doc, bullets_llegadas)
    
    ############
    # Disclaimer
    ############
    # Salto de página
    doc.add_page_break()
    # Agregar saltos de línea para centrar el texto verticalmente
    for _ in range(12): 
        doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
    # Agregar el texto del disclaimer
    paragraph_disclaimer = doc.add_paragraph(disclaimer)
    # Formato del texto
    run = paragraph_disclaimer.runs[0]
    run.font.name = 'Century Gothic'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 32, 96)
    run.bold = True
    paragraph_disclaimer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Guardar el documento
    doc.save(file_path)



def create_document_hubs(tablas, file_path, titulo, header_image_left, footer_image, session, geo_params):
  
    doc = Document()
    estilos(doc)

    # Agregar encabezado y pie de página
    footer_text = """Calle 28 # 13ª - 15, Edificio CCI Pisos 35 - 36 | Bogotá, Colombia T: +57 (1) 560 0100 | info@procolombia.co | www.procolombia.co"""

    add_header_footer(doc, header_image_left, footer_image, footer_text)
        
    # Agregar el título principal del informe
    title_paragraph = doc.add_paragraph(f'TRES EJES HUBS: {str(titulo).upper()}', style='Title')
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Agregar la fecha
    doc_params = obtener_parametros_documento(session)
    fecha = doc_params['Fecha de actualización']
    date_paragraph = doc.add_paragraph(f'ÚLTIMA ACTUALIZACIÓN: {fecha.upper()}', style='Title')
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
     
    # Agregar la tabla de contenido
    agregar_tabla_contenidos(doc, 10)
    doc.add_page_break()

    # Fuentes del documento:
    fuente_exportaciones = 'DANE-DIAN. Cálculos: ProColombia.'
    fuente_inversion = 'Banco de la República. Cálculos: ProColombia.'
    # Fuente condicional de Turismo con Venezuela
    if '850' in geo_params['PAISES_TURISMO_COD']:
        fuente_turismo = 'Migración Colombia: ProColombia. Nota: Los datos excluyen el registro de residentes venezolanos reportado por Migración Colombia (sin incluir la estimación del MinCIT). Tampoco se incluyen las cifras de colombianos residentes en el exterior ni de cruceristas.'
    else:
        fuente_turismo = 'Migración Colombia: ProColombia. Nota: Los datos no incluyen las cifras de colombianos residentes en el exterior ni de cruceristas.'

    # Disclaimer:
    disclaimer = "La información contenida en este documento es de orientación y guía general. En ningún caso, ProColombia, ni sus empleados, son responsables ante usted o cualquier otra persona por las decisiones o acciones que pueda tomar en relación con la información proporcionada, por lo cual debe tomarse como de carácter referencial únicamente."

    # Obtener parámetos de veríficación
    dict_verificacion = verif_ejes(session, geo_params)

    #########
    # RESUMEN
    #########
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO') or (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO') or (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO') or (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
        add_heading(doc, 'Resumen', level=2, style='Heading 1')
        # Exportaciones
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_heading(doc, 'Exportaciones', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_expo"]), 'Table Grid', 10, fuente_exportaciones)

            # Lista para los bullets
            bullets_exportaciones = []

            # Agregar los bullets según las condiciones
            if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
                if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b1_corrido"])
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b2_corrido"])
                else:
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_cerrado"])
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_cerrado"])
                # Siempre se incluye el bullet 3 de empresas:
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b3_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b3_corrido"])
            elif (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_corrido"])
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_corrido"])
                # Siempre se incluye el bullet 3 de empresas:
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b3_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b3_corrido"])

            # Usar la función para agregar los bullets al documento
            add_bullet_points(doc, bullets_exportaciones)

        # Inversión
        if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO') or (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
            add_heading(doc, 'Inversión', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_inv"]), 'Table Grid', 10, fuente_inversion)
            
            # Lista para los bullets de inversión
            bullets_inversion = []

            # Agregar los bullets de inversión según las condiciones
            if dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO':
                if dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO':
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b1_cerrado"] + " " + tablas["RESUMEN"]["texto_inversion_b1_corrido"])
                else:
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b1_cerrado"])
            elif dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO':
                bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b1_corrido"])

            if dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO':
                if dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO':
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b2_cerrado"] + " " + tablas["RESUMEN"]["texto_inversion_b2_corrido"])
                else:
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b2_cerrado"])
            elif dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO':
                bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b2_corrido"])

            # Usar la función para agregar los bullets de inversión al documento
            add_bullet_points(doc, bullets_inversion)


        # Turismo
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_heading(doc, 'Turismo', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_tur"]), 'Table Grid', 10, fuente_turismo)

            # Agregar los bullets de turismo según las condiciones
            if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
                add_bullet_points(doc, [tablas["RESUMEN"]["texto_turismo_b1_cerrado"]])    
            if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
                add_bullet_points(doc, [tablas["RESUMEN"]["texto_turismo_b2_corrido"]])

        # Salto de página
        doc.add_page_break()       

    ###############
    # Exportaciones
    ###############
    add_heading(doc, 'Exportaciones', level=2, style='Heading 1')
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
        
        #####################
        # Tipo de exportación
        #####################
        add_heading(doc, 'Tipo de exportación', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TIPOS"]["ST_CATEGORIAS_CERRADO"]), 'Table Grid', 10, fuente_exportaciones)
        # Año corrido
        if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TIPOS"]["ST_CATEGORIAS_CORRIDO"]), 'Table Grid', 10, fuente_exportaciones)

        #####################################
        # Exportaciones no minero-energéticas
        #####################################
        if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO') or (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
            add_heading(doc, 'Exportaciones no minero-energéticas', level=2, style='Heading 1')

            ###########
            #  Destinos
            ###########
            add_heading(doc, 'Destinos', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'): 
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["PAIS"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["PAIS"]), 'Table Grid', 10, fuente_exportaciones)
            
            ########################
            # Departamento de origen
            ########################
            add_heading(doc, 'Departamento de origen', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["DEPARTAMENTOS"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["DEPARTAMENTOS"]), 'Table Grid', 10, fuente_exportaciones)

            ########
            # Sector
            ########
            add_heading(doc, 'Sector', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["SECTORES"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["SECTORES"]), 'Table Grid', 10, fuente_exportaciones)

            #############
            # SubSectores
            #############
            add_heading(doc, 'Subsector', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["SUBSECTORES"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["SUBSECTORES"]), 'Table Grid', 10, fuente_exportaciones)

            ####################
            # Conteo de Empresas
            ####################
            if (dict_verificacion['exportaciones_conteo_cerrado'] == 'CON DATOS DE CONTEO CERRADO') or (dict_verificacion['exportaciones_conteo_corrido'] == 'CON DATOS DE CONTEO CORRIDO'):
                add_heading(doc, 'Empresas', level=3, style='Heading 2')
                # Año cerrado
                if (dict_verificacion['exportaciones_conteo_cerrado'] == 'CON DATOS DE CONTEO CERRADO'):
                    year_cerrado = doc_params['Año cerrado (T)']
                    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
                    conteo_cerrado = tablas['CONTEO EMPRESAS']['CERRADO'][0]
                    run_cerrado = paragraph_cerrado.add_run(f"{conteo_cerrado} empresas")
                    run_cerrado.bold = True
                elif (dict_verificacion['exportaciones_conteo_cerrado'] == 'SIN DATOS DE CONTEO CERRADO'):
                    year_cerrado = doc_params['Año cerrado (T)']
                    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
                    conteo_cerrado = "0"
                    run_cerrado = paragraph_cerrado.add_run(f"{conteo_cerrado} empresas")
                    run_cerrado.bold = True                    
                # Año corrido
                if (dict_verificacion['exportaciones_conteo_corrido'] == 'CON DATOS DE CONTEO CORRIDO'):
                    year_corrido = doc_params['Año corrido texto (T)']
                    paragraph_corrido = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
                    conteo_corrido = tablas['CONTEO EMPRESAS']['CORRIDO'][0]
                    run_corrido = paragraph_corrido.add_run(f"{conteo_corrido} empresas")
                    run_corrido.bold = True
                elif (dict_verificacion['exportaciones_conteo_corrido'] == 'SIN DATOS DE CONTEO CORRIDO'):
                    year_corrido = doc_params['Año corrido texto (T)']
                    paragraph_corrido = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
                    conteo_corrido = "0"
                    run_corrido = paragraph_corrido.add_run(f"{conteo_corrido} empresas")
                    run_corrido.bold = True

            ################
            # Datos empresas
            ################
            if (dict_verificacion['exportaciones_empresas_cerrado'] == 'CON DATOS DE EMPRESAS CERRADO') or (dict_verificacion['exportaciones_empresas_corrido'] == 'CON DATOS DE EMPRESAS CORRIDO'):
                add_heading(doc, 'Información de Empresas', level=3, style='Heading 2')
                # Año cerrado
                if (dict_verificacion['exportaciones_empresas_cerrado'] == 'CON DATOS DE EMPRESAS CERRADO'):
                    add_table(doc, pd.DataFrame(tablas['EMPRESAS']['ST_NIT_CERRADO']), 'Table Grid', 9, fuente_exportaciones)
                # Año corrido
                if (dict_verificacion['exportaciones_empresas_corrido'] == 'CON DATOS DE EMPRESAS CORRIDO'):
                    add_table(doc, pd.DataFrame(tablas['EMPRESAS']['ST_NIT_CORRIDO']), 'Table Grid', 9, fuente_exportaciones)

            ##############################
            # Oportunidades de exportación
            ##############################
            if (dict_verificacion['oportunidades_exportacion'] == "CON OPORTUNIDADES"):
                add_heading(doc, 'Oportunidades de exportación identificadas', level=3, style='Heading 2')
                agregar_oportunidades_al_documento(doc, tablas, 'OPORTUNIDADES_EXPORTACIONES')

        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('El HUB no registra datos de exportaciones.')

    ###########
    # Inversión
    ###########
    add_heading(doc, 'Inversión', level=2, style='Heading 1')
    if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO') or (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
        
        #####
        # IED
        #####
        # Total
        if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO'):
            add_heading(doc, 'IED', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO'):
                add_table(doc, pd.DataFrame(tablas['IED TOTAL']['ied_cerrado_total']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['IED TOTAL']['ied_corrido_total']), 'Table Grid', 10, fuente_inversion)
            # Países
            add_heading(doc, 'IED - Países', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO'):
                add_table(doc, pd.DataFrame(tablas['IED PAISES']['ied_cerrado']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['IED PAISES']['ied_corrido']), 'Table Grid', 10, fuente_inversion)

        #####
        # ICE
        #####
        # Total
        if (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
            add_heading(doc, 'ICE', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO'):
                add_table(doc, pd.DataFrame(tablas['ICE TOTAL']['ice_cerrado_total']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['ICE TOTAL']['ice_corrido_total']), 'Table Grid', 10, fuente_inversion)
            # Países
            add_heading(doc, 'ICE - Países', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO'):
                add_table(doc, pd.DataFrame(tablas['ICE PAISES']['ice_cerrado']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['ICE PAISES']['ice_corrido']), 'Table Grid', 10, fuente_inversion)

        ############################
        # Oportunidades de inversión
        ############################
        if (dict_verificacion['oportunidades_inversion'] == "CON OPORTUNIDADES"):
            add_heading(doc, 'Oportunidades de inversión identificadas', level=3, style='Heading 2')
            agregar_oportunidades_al_documento(doc, tablas, 'OPORTUNIDADES_INVERSION')
            
        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('El HUB no registra datos de inversión.')

    #########
    # Turismo
    #########
    add_heading(doc, 'Turismo', level=2, style='Heading 1')
    if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
        
        ########
        # Países
        ########
        add_heading(doc, 'Países', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["PAIS_RESIDENCIA"]), 'Table Grid', 10, fuente_turismo)
        # Año corrido
        if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CORRIDO"]["PAIS_RESIDENCIA"]), 'Table Grid', 10, fuente_turismo)

        ###############
        # Departamentos
        ###############
        add_heading(doc, 'Departamentos', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["DPTO_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)
        # Año corrido
        if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CORRIDO"]["DPTO_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)

        ############
        # Municipios
        ############
        add_heading(doc, 'Municipios', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["CIUDAD_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)
        # Año corrido
        if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CORRIDO"]["CIUDAD_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)

        ########
        # Género
        ########
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_heading(doc, 'Género', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["DESCRIPCION_GENERO"]), 'Table Grid', 10, fuente_turismo)

        #########
        # Motivos
        #########
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_heading(doc, 'Motivos', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["MOVC_NOMBRE"]), 'Table Grid', 10, fuente_turismo)

        ###############
        # OPORTUNIDADES
        ###############
        if (dict_verificacion['oportunidades_turismo'] == "CON OPORTUNIDADES"):
            add_heading(doc, 'Productos vacacionales que se promocionan', level=3, style='Heading 2')
            # Principales
            p1 = doc.add_paragraph()
            run1 = p1.add_run("Principales: ")
            run1.bold = True
            p1.add_run(tablas['TURISMO_PRINCIPAL'])
            p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Nichos
            p2 = doc.add_paragraph()
            run2 = p2.add_run("Nichos: ")
            run2.bold = True
            p2.add_run(tablas['TURISMO_NICHOS'])
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('El HUB no registra datos de turismo.')

    ###########
    # Logística
    ###########
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
        add_heading(doc, 'Logística', level=2, style='Heading 1')
    #######
    # Pesos
    #######
        add_heading(doc, 'Pesos', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
            add_table(doc, pd.DataFrame(tablas['TIPOS PESO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
        # Año corrido
        if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_table(doc, pd.DataFrame(tablas['TIPOS PESO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)
        # Mineros
        if (dict_verificacion['pesos_minero_cerrado'] == 'CON DATOS CERRADO') or (dict_verificacion['pesos_minero_corrido'] == 'CON DATOS CORRIDO'):
            add_heading(doc, 'Pesos por medio de transporte: exportaciones mineras', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO MINERO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO MINERO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)
        # No Mineros
        if (dict_verificacion['pesos_no_minero_cerrado'] == 'CON DATOS CERRADO') or (dict_verificacion['pesos_no_minero_corrido'] == 'CON DATOS CORRIDO'):
            add_heading(doc, 'Pesos por medio de transporte: exportaciones no mineras', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO NO MINERO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO NO MINERO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)
        # Salto de página
        doc.add_page_break()
   
    #######
    # Anexo
    #######
    add_heading(doc, 'Anexo: Países considerados', level=2, style='Heading 1')
    # Agregar el párrafo introductorio
    paragraph_intro = doc.add_paragraph(
        f'El presente documento muestra los datos agregados para los Tres Ejes de negocio de ProColombia para el continente de {str(titulo).capitalize()}. Se incluye información de los siguientes países:')
    paragraph_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Agregar el párrafo con los países, añadir un punto al final y justificarlo
    paises_anexo_con_punto = geo_params['PAISES_ANEXO'] + '.'
    paragraph_paises = doc.add_paragraph(paises_anexo_con_punto)
    paragraph_paises.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    ############
    # Disclaimer
    ############
    # Salto de página
    doc.add_page_break()
    # Agregar saltos de línea para centrar el texto verticalmente
    for _ in range(12): 
        doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
    # Agregar el texto del disclaimer
    paragraph_disclaimer = doc.add_paragraph(disclaimer)
    # Formato del texto
    run = paragraph_disclaimer.runs[0]
    run.font.name = 'Century Gothic'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 32, 96)
    run.bold = True
    paragraph_disclaimer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Guardar el documento
    doc.save(file_path)


def create_document_tlcs(tablas, file_path, titulo, header_image_left, footer_image, session, geo_params):
  
    doc = Document()
    estilos(doc)

    # Agregar encabezado y pie de página
    footer_text = """Calle 28 # 13ª - 15, Edificio CCI Pisos 35 - 36 | Bogotá, Colombia T: +57 (1) 560 0100 | info@procolombia.co | www.procolombia.co"""

    add_header_footer(doc, header_image_left, footer_image, footer_text)
        
    # Agregar el título principal del informe
    title_paragraph = doc.add_paragraph(f'TRES EJES TLCS: {str(titulo).upper()}', style='Title')
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Agregar la fecha
    doc_params = obtener_parametros_documento(session)
    fecha = doc_params['Fecha de actualización']
    date_paragraph = doc.add_paragraph(f'ÚLTIMA ACTUALIZACIÓN: {fecha.upper()}', style='Title')
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
     
    # Agregar la tabla de contenido
    agregar_tabla_contenidos(doc, 10)
    doc.add_page_break()

    # Fuentes del documento:
    fuente_exportaciones = 'DANE-DIAN. Cálculos: ProColombia.'
    fuente_inversion = 'Banco de la República. Cálculos: ProColombia.'
    # Fuente condicional de Turismo con Venezuela
    if '850' in geo_params['PAISES_TURISMO_COD']:
        fuente_turismo = 'Migración Colombia: ProColombia. Nota: Los datos excluyen el registro de residentes venezolanos reportado por Migración Colombia (sin incluir la estimación del MinCIT). Tampoco se incluyen las cifras de colombianos residentes en el exterior ni de cruceristas.'
    else:
        fuente_turismo = 'Migración Colombia: ProColombia. Nota: Los datos no incluyen las cifras de colombianos residentes en el exterior ni de cruceristas.'

    # Disclaimer:
    disclaimer = "La información contenida en este documento es de orientación y guía general. En ningún caso, ProColombia, ni sus empleados, son responsables ante usted o cualquier otra persona por las decisiones o acciones que pueda tomar en relación con la información proporcionada, por lo cual debe tomarse como de carácter referencial únicamente."

    # Obtener parámetos de veríficación
    dict_verificacion = verif_ejes(session, geo_params)

    #########
    # RESUMEN
    #########
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO') or (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO') or (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO') or (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
        add_heading(doc, 'Resumen', level=2, style='Heading 1')
        # Exportaciones
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_heading(doc, 'Exportaciones', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_expo"]), 'Table Grid', 10, fuente_exportaciones)

            # Lista para los bullets
            bullets_exportaciones = []

            # Agregar los bullets según las condiciones
            if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
                if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b1_corrido"])
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b2_corrido"])
                else:
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_cerrado"])
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_cerrado"])
                # Siempre se incluye el bullet 3 de empresas:
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b3_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b3_corrido"])
            elif (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_corrido"])
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_corrido"])
                # Siempre se incluye el bullet 3 de empresas:
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b3_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b3_corrido"])

            # Usar la función para agregar los bullets al documento
            add_bullet_points(doc, bullets_exportaciones)

        # Inversión
        if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO') or (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
            add_heading(doc, 'Inversión', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_inv"]), 'Table Grid', 10, fuente_inversion)
            
            # Lista para los bullets de inversión
            bullets_inversion = []

            # Agregar los bullets de inversión según las condiciones
            if dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO':
                if dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO':
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b1_cerrado"] + " " + tablas["RESUMEN"]["texto_inversion_b1_corrido"])
                else:
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b1_cerrado"])
            elif dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO':
                bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b1_corrido"])

            if dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO':
                if dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO':
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b2_cerrado"] + " " + tablas["RESUMEN"]["texto_inversion_b2_corrido"])
                else:
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b2_cerrado"])
            elif dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO':
                bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b2_corrido"])

            # Usar la función para agregar los bullets de inversión al documento
            add_bullet_points(doc, bullets_inversion)


        # Turismo
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_heading(doc, 'Turismo', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_tur"]), 'Table Grid', 10, fuente_turismo)

            # Agregar los bullets de turismo según las condiciones
            if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
                add_bullet_points(doc, [tablas["RESUMEN"]["texto_turismo_b1_cerrado"]])    
            if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
                add_bullet_points(doc, [tablas["RESUMEN"]["texto_turismo_b2_corrido"]])

        # Salto de página
        doc.add_page_break()

    ###############
    # Exportaciones
    ###############
    add_heading(doc, 'Exportaciones', level=2, style='Heading 1')
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
        
        #####################
        # Tipo de exportación
        #####################
        add_heading(doc, 'Tipo de exportación', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TIPOS"]["ST_CATEGORIAS_CERRADO"]), 'Table Grid', 10, fuente_exportaciones)
        # Año corrido
        if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TIPOS"]["ST_CATEGORIAS_CORRIDO"]), 'Table Grid', 10, fuente_exportaciones)

        #####################################
        # Exportaciones no minero-energéticas
        #####################################
        if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO') or (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
            add_heading(doc, 'Exportaciones no minero-energéticas', level=2, style='Heading 1')

            ###########
            #  Destinos
            ###########
            add_heading(doc, 'Destinos', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'): 
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["PAIS"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["PAIS"]), 'Table Grid', 10, fuente_exportaciones)
            
            ########################
            # Departamento de origen
            ########################
            add_heading(doc, 'Departamento de origen', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["DEPARTAMENTOS"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["DEPARTAMENTOS"]), 'Table Grid', 10, fuente_exportaciones)

            ########
            # Sector
            ########
            add_heading(doc, 'Sector', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["SECTORES"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["SECTORES"]), 'Table Grid', 10, fuente_exportaciones)

            #############
            # SubSectores
            #############
            add_heading(doc, 'Subsector', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["SUBSECTORES"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["SUBSECTORES"]), 'Table Grid', 10, fuente_exportaciones)

            ####################
            # Conteo de Empresas
            ####################
            if (dict_verificacion['exportaciones_conteo_cerrado'] == 'CON DATOS DE CONTEO CERRADO') or (dict_verificacion['exportaciones_conteo_corrido'] == 'CON DATOS DE CONTEO CORRIDO'):
                add_heading(doc, 'Empresas', level=3, style='Heading 2')
                # Año cerrado
                if (dict_verificacion['exportaciones_conteo_cerrado'] == 'CON DATOS DE CONTEO CERRADO'):
                    year_cerrado = doc_params['Año cerrado (T)']
                    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
                    conteo_cerrado = tablas['CONTEO EMPRESAS']['CERRADO'][0]
                    run_cerrado = paragraph_cerrado.add_run(f"{conteo_cerrado} empresas")
                    run_cerrado.bold = True
                elif (dict_verificacion['exportaciones_conteo_cerrado'] == 'SIN DATOS DE CONTEO CERRADO'):
                    year_cerrado = doc_params['Año cerrado (T)']
                    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
                    conteo_cerrado = "0"
                    run_cerrado = paragraph_cerrado.add_run(f"{conteo_cerrado} empresas")
                    run_cerrado.bold = True                    
                # Año corrido
                if (dict_verificacion['exportaciones_conteo_corrido'] == 'CON DATOS DE CONTEO CORRIDO'):
                    year_corrido = doc_params['Año corrido texto (T)']
                    paragraph_corrido = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
                    conteo_corrido = tablas['CONTEO EMPRESAS']['CORRIDO'][0]
                    run_corrido = paragraph_corrido.add_run(f"{conteo_corrido} empresas")
                    run_corrido.bold = True
                elif (dict_verificacion['exportaciones_conteo_corrido'] == 'SIN DATOS DE CONTEO CORRIDO'):
                    year_corrido = doc_params['Año corrido texto (T)']
                    paragraph_corrido = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
                    conteo_corrido = "0"
                    run_corrido = paragraph_corrido.add_run(f"{conteo_corrido} empresas")
                    run_corrido.bold = True

            ################
            # Datos empresas
            ################
            if (dict_verificacion['exportaciones_empresas_cerrado'] == 'CON DATOS DE EMPRESAS CERRADO') or (dict_verificacion['exportaciones_empresas_corrido'] == 'CON DATOS DE EMPRESAS CORRIDO'):
                add_heading(doc, 'Información de Empresas', level=3, style='Heading 2')
                # Año cerrado
                if (dict_verificacion['exportaciones_empresas_cerrado'] == 'CON DATOS DE EMPRESAS CERRADO'):
                    add_table(doc, pd.DataFrame(tablas['EMPRESAS']['ST_NIT_CERRADO']), 'Table Grid', 9, fuente_exportaciones)
                # Año corrido
                if (dict_verificacion['exportaciones_empresas_corrido'] == 'CON DATOS DE EMPRESAS CORRIDO'):
                    add_table(doc, pd.DataFrame(tablas['EMPRESAS']['ST_NIT_CORRIDO']), 'Table Grid', 9, fuente_exportaciones)

            ##############################
            # Oportunidades de exportación
            ##############################
            if (dict_verificacion['oportunidades_exportacion'] == "CON OPORTUNIDADES"):
                add_heading(doc, 'Oportunidades de exportación identificadas', level=3, style='Heading 2')
                agregar_oportunidades_al_documento(doc, tablas, 'OPORTUNIDADES_EXPORTACIONES')
        
        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('El TLC no registra datos de exportaciones.')

    ###########
    # Inversión
    ###########
    add_heading(doc, 'Inversión', level=2, style='Heading 1')
    if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO') or (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
        
        #####
        # IED
        #####
        # Total
        if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO'):
            add_heading(doc, 'IED', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO'):
                add_table(doc, pd.DataFrame(tablas['IED TOTAL']['ied_cerrado_total']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['IED TOTAL']['ied_corrido_total']), 'Table Grid', 10, fuente_inversion)
            # Países
            add_heading(doc, 'IED - Países', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO'):
                add_table(doc, pd.DataFrame(tablas['IED PAISES']['ied_cerrado']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['IED PAISES']['ied_corrido']), 'Table Grid', 10, fuente_inversion)

        #####
        # ICE
        #####
        # Total
        if (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
            add_heading(doc, 'ICE', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO'):
                add_table(doc, pd.DataFrame(tablas['ICE TOTAL']['ice_cerrado_total']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['ICE TOTAL']['ice_corrido_total']), 'Table Grid', 10, fuente_inversion)
            # Países
            add_heading(doc, 'ICE - Países', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO'):
                add_table(doc, pd.DataFrame(tablas['ICE PAISES']['ice_cerrado']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['ICE PAISES']['ice_corrido']), 'Table Grid', 10, fuente_inversion)

        ############################
        # Oportunidades de inversión
        ############################
        if (dict_verificacion['oportunidades_inversion'] == "CON OPORTUNIDADES"):
            add_heading(doc, 'Oportunidades de inversión identificadas', level=3, style='Heading 2')
            agregar_oportunidades_al_documento(doc, tablas, 'OPORTUNIDADES_INVERSION')
            
        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('El TLC no registra datos de inversión.')
    
    #########
    # Turismo
    #########
    add_heading(doc, 'Turismo', level=2, style='Heading 1')
    if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
        
        ########
        # Países
        ########
        add_heading(doc, 'Países', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["PAIS_RESIDENCIA"]), 'Table Grid', 10, fuente_turismo)
        # Año corrido
        if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CORRIDO"]["PAIS_RESIDENCIA"]), 'Table Grid', 10, fuente_turismo)

        ###############
        # Departamentos
        ###############
        add_heading(doc, 'Departamentos', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["DPTO_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)
        # Año corrido
        if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CORRIDO"]["DPTO_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)

        ############
        # Municipios
        ############
        add_heading(doc, 'Municipios', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["CIUDAD_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)
        # Año corrido
        if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CORRIDO"]["CIUDAD_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)

        ########
        # Género
        ########
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_heading(doc, 'Género', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["DESCRIPCION_GENERO"]), 'Table Grid', 10, fuente_turismo)

        #########
        # Motivos
        #########
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_heading(doc, 'Motivos', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["MOVC_NOMBRE"]), 'Table Grid', 10, fuente_turismo)

        ###############
        # OPORTUNIDADES
        ###############
        if (dict_verificacion['oportunidades_turismo'] == "CON OPORTUNIDADES"):
            add_heading(doc, 'Productos vacacionales que se promocionan', level=3, style='Heading 2')
            # Principales
            p1 = doc.add_paragraph()
            run1 = p1.add_run("Principales: ")
            run1.bold = True
            p1.add_run(tablas['TURISMO_PRINCIPAL'])
            p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Nichos
            p2 = doc.add_paragraph()
            run2 = p2.add_run("Nichos: ")
            run2.bold = True
            p2.add_run(tablas['TURISMO_NICHOS'])
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('El TLC no registra datos de turismo.')

    ###########
    # Logística
    ###########
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
        add_heading(doc, 'Logística', level=2, style='Heading 1')
    #######
    # Pesos
    #######
        add_heading(doc, 'Pesos', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
            add_table(doc, pd.DataFrame(tablas['TIPOS PESO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
        # Año corrido
        if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_table(doc, pd.DataFrame(tablas['TIPOS PESO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)
        # Mineros
        if (dict_verificacion['pesos_minero_cerrado'] == 'CON DATOS CERRADO') or (dict_verificacion['pesos_minero_corrido'] == 'CON DATOS CORRIDO'):
            add_heading(doc, 'Pesos por medio de transporte: exportaciones mineras', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO MINERO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO MINERO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)
        # No Mineros
        if (dict_verificacion['pesos_no_minero_cerrado'] == 'CON DATOS CERRADO') or (dict_verificacion['pesos_no_minero_corrido'] == 'CON DATOS CORRIDO'):
            add_heading(doc, 'Pesos por medio de transporte: exportaciones no mineras', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO NO MINERO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO NO MINERO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)
        # Salto de página
        doc.add_page_break()
        
    #######
    # Anexo
    #######
    add_heading(doc, 'Anexo: Países considerados', level=2, style='Heading 1')
    # Agregar el párrafo introductorio
    paragraph_intro = doc.add_paragraph(
        f'El presente documento muestra los datos agregados para los Tres Ejes de negocio de ProColombia para el continente de {str(titulo).capitalize()}. Se incluye información de los siguientes países:')
    paragraph_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Agregar el párrafo con los países, añadir un punto al final y justificarlo
    paises_anexo_con_punto = geo_params['PAISES_ANEXO'] + '.'
    paragraph_paises = doc.add_paragraph(paises_anexo_con_punto)
    paragraph_paises.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    ############
    # Disclaimer
    ############
    # Salto de página
    doc.add_page_break()
    # Agregar saltos de línea para centrar el texto verticalmente
    for _ in range(12): 
        doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
    # Agregar el texto del disclaimer
    paragraph_disclaimer = doc.add_paragraph(disclaimer)
    # Formato del texto
    run = paragraph_disclaimer.runs[0]
    run.font.name = 'Century Gothic'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 32, 96)
    run.bold = True
    paragraph_disclaimer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Guardar el documento
    doc.save(file_path)


def create_document_paises(tablas, file_path, titulo, header_image_left, footer_image, session, geo_params):
  
    doc = Document()
    estilos(doc)

    # Agregar encabezado y pie de página
    footer_text = """Calle 28 # 13ª - 15, Edificio CCI Pisos 35 - 36 | Bogotá, Colombia T: +57 (1) 560 0100 | info@procolombia.co | www.procolombia.co"""

    add_header_footer(doc, header_image_left, footer_image, footer_text)
        
    # Agregar el título principal del informe
    title_paragraph = doc.add_paragraph(f'TRES EJES PAÍSES: {str(titulo).upper()}', style='Title')
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Agregar la fecha
    doc_params = obtener_parametros_documento(session)
    fecha = doc_params['Fecha de actualización']
    date_paragraph = doc.add_paragraph(f'ÚLTIMA ACTUALIZACIÓN: {fecha.upper()}', style='Title')
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
     
    # Agregar la tabla de contenido
    agregar_tabla_contenidos(doc, 10)
    doc.add_page_break()

    # Fuentes del documento:
    fuente_exportaciones = 'DANE-DIAN. Cálculos: ProColombia.'
    fuente_inversion = 'Banco de la República. Cálculos: ProColombia.'
    # Fuente condicional de Turismo con Venezuela
    if '850' in geo_params['PAISES_TURISMO_COD']:
        fuente_turismo = 'Migración Colombia: ProColombia. Nota: Los datos excluyen el registro de residentes venezolanos reportado por Migración Colombia (sin incluir la estimación del MinCIT). Tampoco se incluyen las cifras de colombianos residentes en el exterior ni de cruceristas.'
    else:
        fuente_turismo = 'Migración Colombia: ProColombia. Nota: Los datos no incluyen las cifras de colombianos residentes en el exterior ni de cruceristas.'

    # Disclaimer:
    disclaimer = "La información contenida en este documento es de orientación y guía general. En ningún caso, ProColombia, ni sus empleados, son responsables ante usted o cualquier otra persona por las decisiones o acciones que pueda tomar en relación con la información proporcionada, por lo cual debe tomarse como de carácter referencial únicamente."

    # Obtener parámetos de veríficación
    dict_verificacion = verif_ejes(session, geo_params)

    #########
    # RESUMEN
    #########
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO') or (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO') or (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO') or (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
        add_heading(doc, 'Resumen', level=2, style='Heading 1')
        # Exportaciones
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_heading(doc, 'Exportaciones', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_expo"]), 'Table Grid', 10, fuente_exportaciones)

            # Lista para los bullets
            bullets_exportaciones = []

            # Agregar los bullets según las condiciones
            if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
                if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b1_corrido"])
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b2_corrido"])
                else:
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_cerrado"])
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_cerrado"])
                # Siempre se incluye el bullet 3 de empresas:
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b3_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b3_corrido"])
            elif (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_corrido"])
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_corrido"])
                # Siempre se incluye el bullet 3 de empresas:
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b3_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b3_corrido"])

            # Usar la función para agregar los bullets al documento
            add_bullet_points(doc, bullets_exportaciones)

        # Inversión
        if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO') or (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
            add_heading(doc, 'Inversión', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_inv"]), 'Table Grid', 10, fuente_inversion)
            
            # Lista para los bullets de inversión
            bullets_inversion = []

            # Agregar los bullets de inversión según las condiciones
            if dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO':
                if dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO':
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b1_cerrado"] + " " + tablas["RESUMEN"]["texto_inversion_b1_corrido"])
                else:
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b1_cerrado"])
            elif dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO':
                bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b1_corrido"])

            if dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO':
                if dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO':
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b2_cerrado"] + " " + tablas["RESUMEN"]["texto_inversion_b2_corrido"])
                else:
                    bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b2_cerrado"])
            elif dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO':
                bullets_inversion.append(tablas["RESUMEN"]["texto_inversion_b2_corrido"])

            # Usar la función para agregar los bullets de inversión al documento
            add_bullet_points(doc, bullets_inversion)


        # Turismo
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_heading(doc, 'Turismo', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_tur"]), 'Table Grid', 10, fuente_turismo)

            # Agregar los bullets de turismo según las condiciones
            if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
                add_bullet_points(doc, [tablas["RESUMEN"]["texto_turismo_b1_cerrado"]])    
            if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
                add_bullet_points(doc, [tablas["RESUMEN"]["texto_turismo_b2_corrido"]])

        # Salto de página
        doc.add_page_break()

    ###############
    # Exportaciones
    ###############
    add_heading(doc, 'Exportaciones', level=2, style='Heading 1')
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
        
        #####################
        # Tipo de exportación
        #####################
        add_heading(doc, 'Tipo de exportación', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TIPOS"]["ST_CATEGORIAS_CERRADO"]), 'Table Grid', 10, fuente_exportaciones)
        # Año corrido
        if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TIPOS"]["ST_CATEGORIAS_CORRIDO"]), 'Table Grid', 10, fuente_exportaciones)

        #####################################
        # Exportaciones no minero-energéticas
        #####################################
        if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO') or (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
            add_heading(doc, 'Exportaciones no minero-energéticas', level=2, style='Heading 1')
            
            ########################
            # Departamento de origen
            ########################
            add_heading(doc, 'Departamento de origen', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["DEPARTAMENTOS"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["DEPARTAMENTOS"]), 'Table Grid', 10, fuente_exportaciones)

            ########
            # Sector
            ########
            add_heading(doc, 'Sector', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["SECTORES"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["SECTORES"]), 'Table Grid', 10, fuente_exportaciones)

            #############
            # SubSectores
            #############
            add_heading(doc, 'Subsector', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["SUBSECTORES"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["SUBSECTORES"]), 'Table Grid', 10, fuente_exportaciones)

            ####################
            # Conteo de Empresas
            ####################
            if (dict_verificacion['exportaciones_conteo_cerrado'] == 'CON DATOS DE CONTEO CERRADO') or (dict_verificacion['exportaciones_conteo_corrido'] == 'CON DATOS DE CONTEO CORRIDO'):
                add_heading(doc, 'Empresas', level=3, style='Heading 2')
                # Año cerrado
                if (dict_verificacion['exportaciones_conteo_cerrado'] == 'CON DATOS DE CONTEO CERRADO'):
                    year_cerrado = doc_params['Año cerrado (T)']
                    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
                    conteo_cerrado = tablas['CONTEO EMPRESAS']['CERRADO'][0]
                    run_cerrado = paragraph_cerrado.add_run(f"{conteo_cerrado} empresas")
                    run_cerrado.bold = True
                elif (dict_verificacion['exportaciones_conteo_cerrado'] == 'SIN DATOS DE CONTEO CERRADO'):
                    year_cerrado = doc_params['Año cerrado (T)']
                    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
                    conteo_cerrado = "0"
                    run_cerrado = paragraph_cerrado.add_run(f"{conteo_cerrado} empresas")
                    run_cerrado.bold = True                    
                # Año corrido
                if (dict_verificacion['exportaciones_conteo_corrido'] == 'CON DATOS DE CONTEO CORRIDO'):
                    year_corrido = doc_params['Año corrido texto (T)']
                    paragraph_corrido = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
                    conteo_corrido = tablas['CONTEO EMPRESAS']['CORRIDO'][0]
                    run_corrido = paragraph_corrido.add_run(f"{conteo_corrido} empresas")
                    run_corrido.bold = True
                elif (dict_verificacion['exportaciones_conteo_corrido'] == 'SIN DATOS DE CONTEO CORRIDO'):
                    year_corrido = doc_params['Año corrido texto (T)']
                    paragraph_corrido = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
                    conteo_corrido = "0"
                    run_corrido = paragraph_corrido.add_run(f"{conteo_corrido} empresas")
                    run_corrido.bold = True

            ################
            # Datos empresas
            ################
            if (dict_verificacion['exportaciones_empresas_cerrado'] == 'CON DATOS DE EMPRESAS CERRADO') or (dict_verificacion['exportaciones_empresas_corrido'] == 'CON DATOS DE EMPRESAS CORRIDO'):
                add_heading(doc, 'Información de Empresas', level=3, style='Heading 2')
                # Año cerrado
                if (dict_verificacion['exportaciones_empresas_cerrado'] == 'CON DATOS DE EMPRESAS CERRADO'):
                    add_table(doc, pd.DataFrame(tablas['EMPRESAS']['ST_NIT_CERRADO']), 'Table Grid', 9, fuente_exportaciones)
                # Año corrido
                if (dict_verificacion['exportaciones_empresas_corrido'] == 'CON DATOS DE EMPRESAS CORRIDO'):
                    add_table(doc, pd.DataFrame(tablas['EMPRESAS']['ST_NIT_CORRIDO']), 'Table Grid', 9, fuente_exportaciones)
            
            ##############################
            # Oportunidades de exportación
            ##############################
            if (dict_verificacion['oportunidades_exportacion'] == "CON OPORTUNIDADES"):
                add_heading(doc, 'Oportunidades de exportación identificadas', level=3, style='Heading 2')
                agregar_oportunidades_al_documento(doc, tablas, 'OPORTUNIDADES_EXPORTACIONES')
    
        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('El país no registra datos de exportaciones.')

    ###########
    # Inversión
    ###########
    add_heading(doc, 'Inversión', level=2, style='Heading 1')
    if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO') or (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
        
        #####
        # IED
        #####
        # Total
        if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO'):
            add_heading(doc, 'IED', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO'):
                add_table(doc, pd.DataFrame(tablas['IED TOTAL']['ied_cerrado_total']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['IED TOTAL']['ied_corrido_total']), 'Table Grid', 10, fuente_inversion)

        #####
        # ICE
        #####
        # Total
        if (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
            add_heading(doc, 'ICE', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO'):
                add_table(doc, pd.DataFrame(tablas['ICE TOTAL']['ice_cerrado_total']), 'Table Grid', 10, fuente_inversion)
            # Año corrido
            if (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO'):
                add_table(doc, pd.DataFrame(tablas['ICE TOTAL']['ice_corrido_total']), 'Table Grid', 10, fuente_inversion)

        ############################
        # Oportunidades de inversión
        ############################
        if (dict_verificacion['oportunidades_inversion'] == "CON OPORTUNIDADES"):
            add_heading(doc, 'Oportunidades de inversión identificadas', level=3, style='Heading 2')
            agregar_oportunidades_al_documento(doc, tablas, 'OPORTUNIDADES_INVERSION')
        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('El país no registra datos de inversión.')
    
    #########
    # Turismo
    #########
    add_heading(doc, 'Turismo', level=2, style='Heading 1')
    if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
        
        ###############
        # Departamentos
        ###############
        add_heading(doc, 'Departamentos', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["DPTO_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)
        # Año corrido
        if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CORRIDO"]["DPTO_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)

        ############
        # Municipios
        ############
        add_heading(doc, 'Municipios', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["CIUDAD_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)
        # Año corrido
        if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CORRIDO"]["CIUDAD_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)

        ########
        # Género
        ########
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_heading(doc, 'Género', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["DESCRIPCION_GENERO"]), 'Table Grid', 10, fuente_turismo)

        #########
        # Motivos
        #########
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_heading(doc, 'Motivos', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["MOVC_NOMBRE"]), 'Table Grid', 10, fuente_turismo)
        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('El país no registra datos de turismo.')

    ###########
    # Logística
    ###########
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
        add_heading(doc, 'Logística', level=2, style='Heading 1')
    #######
    # Pesos
    #######
        add_heading(doc, 'Pesos', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
            add_table(doc, pd.DataFrame(tablas['TIPOS PESO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
        # Año corrido
        if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_table(doc, pd.DataFrame(tablas['TIPOS PESO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)
        # Mineros
        if (dict_verificacion['pesos_minero_cerrado'] == 'CON DATOS CERRADO') or (dict_verificacion['pesos_minero_corrido'] == 'CON DATOS CORRIDO'):
            add_heading(doc, 'Pesos por medio de transporte: exportaciones mineras', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO MINERO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO MINERO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)
        # No Mineros
        if (dict_verificacion['pesos_no_minero_cerrado'] == 'CON DATOS CERRADO') or (dict_verificacion['pesos_no_minero_corrido'] == 'CON DATOS CORRIDO'):
            add_heading(doc, 'Pesos por medio de transporte: exportaciones no mineras', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO NO MINERO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO NO MINERO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)

    ############
    # Disclaimer
    ############
    # Salto de página
    doc.add_page_break()
    # Agregar saltos de línea para centrar el texto verticalmente
    for _ in range(12): 
        doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
    # Agregar el texto del disclaimer
    paragraph_disclaimer = doc.add_paragraph(disclaimer)
    # Formato del texto
    run = paragraph_disclaimer.runs[0]
    run.font.name = 'Century Gothic'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 32, 96)
    run.bold = True
    paragraph_disclaimer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Guardar el documento
    doc.save(file_path)


def create_document_departamentos(tablas, file_path, titulo, header_image_left, footer_image, session, geo_params):
  
    doc = Document()
    estilos(doc)

    # Agregar encabezado y pie de página
    footer_text = """Calle 28 # 13ª - 15, Edificio CCI Pisos 35 - 36 | Bogotá, Colombia T: +57 (1) 560 0100 | info@procolombia.co | www.procolombia.co"""

    add_header_footer(doc, header_image_left, footer_image, footer_text)
        
    # Agregar el título principal del informe
    title_paragraph = doc.add_paragraph(f'TRES EJES DEPARTAMENTOS: {str(titulo).upper()}', style='Title')
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Agregar la fecha
    doc_params = obtener_parametros_documento(session)
    fecha = doc_params['Fecha de actualización']
    date_paragraph = doc.add_paragraph(f'ÚLTIMA ACTUALIZACIÓN: {fecha.upper()}', style='Title')
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
     
    # Agregar la tabla de contenido
    agregar_tabla_contenidos(doc, 10)
    doc.add_page_break()

    # Fuentes del documento:
    fuente_exportaciones = 'DANE-DIAN. Cálculos: ProColombia.'
    fuente_inversion = 'Banco de la República. Cálculos: ProColombia.'
    fuente_turismo = 'Migración Colombia: ProColombia. Nota: Los datos excluyen el registro de residentes venezolanos reportado por Migración Colombia (sin incluir la estimación del MinCIT). Tampoco se incluyen las cifras de colombianos residentes en el exterior ni de cruceristas.'
    fuente_conectividad = 'OAG. Cálculos: ProColombia.'

    # Disclaimer:
    disclaimer = "La información contenida en este documento es de orientación y guía general. En ningún caso, ProColombia, ni sus empleados, son responsables ante usted o cualquier otra persona por las decisiones o acciones que pueda tomar en relación con la información proporcionada, por lo cual debe tomarse como de carácter referencial únicamente."

    # Obtener parámetos de veríficación
    dict_verificacion = verif_ejes(session, geo_params)

    #########
    # RESUMEN
    #########
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO') or (dict_verificacion['ied_cerrado'] == 'CON DATOS DE IED CERRADO') or (dict_verificacion['ied_corrido'] == 'CON DATOS DE IED CORRIDO') or (dict_verificacion['ice_cerrado'] == 'CON DATOS DE ICE CERRADO') or (dict_verificacion['ice_corrido'] == 'CON DATOS DE ICE CORRIDO') or (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
        add_heading(doc, 'Resumen', level=2, style='Heading 1')
        # Exportaciones
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_heading(doc, 'Exportaciones', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_expo"]), 'Table Grid', 10, fuente_exportaciones)

            # Lista para los bullets
            bullets_exportaciones = []

            # Agregar los bullets según las condiciones
            if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
                if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b1_corrido"])
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b2_corrido"])
                else:
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_cerrado"])
                    bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_cerrado"])
                # Siempre se incluye el bullet 3 de empresas:
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b3_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b3_corrido"])
            elif (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b1_corrido"])
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b2_corrido"])
                # Siempre se incluye el bullet 3 de empresas:
                bullets_exportaciones.append(tablas["RESUMEN"]["texto_exportaciones_b3_cerrado"] + " " + tablas["RESUMEN"]["texto_exportaciones_b3_corrido"])

            # Usar la función para agregar los bullets al documento
            add_bullet_points(doc, bullets_exportaciones)

        # Turismo
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_heading(doc, 'Turismo', level=3, style='Heading 3')
            add_table_resumen(doc, pd.DataFrame(tablas["RESUMEN"]["tab_resumen_tur"]), 'Table Grid', 10, fuente_turismo)

            # Agregar los bullets de turismo según las condiciones
            if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
                add_bullet_points(doc, [tablas["RESUMEN"]["texto_turismo_b1_cerrado"]])    
            if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
                add_bullet_points(doc, [tablas["RESUMEN"]["texto_turismo_b2_corrido"]])

        # Salto de página
        doc.add_page_break()

    ###############
    # Exportaciones
    ###############
    add_heading(doc, 'Exportaciones', level=2, style='Heading 1')
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
        
        #####################
        # Tipo de exportación
        #####################
        add_heading(doc, 'Tipo de exportación', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TIPOS"]["ST_CATEGORIAS_CERRADO"]), 'Table Grid', 10, fuente_exportaciones)
        # Año corrido
        if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TIPOS"]["ST_CATEGORIAS_CORRIDO"]), 'Table Grid', 10, fuente_exportaciones)

        #####################################
        # Exportaciones no minero-energéticas
        #####################################
        if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO') or (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
            add_heading(doc, 'Exportaciones no minero-energéticas', level=2, style='Heading 1')

            ###########
            #  Destinos
            ###########
            add_heading(doc, 'Destinos', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'): 
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["PAIS"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["PAIS"]), 'Table Grid', 10, fuente_exportaciones)

            ########
            # Sector
            ########
            add_heading(doc, 'Sector', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["SECTORES"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["SECTORES"]), 'Table Grid', 10, fuente_exportaciones)

            #############
            # SubSectores
            #############
            add_heading(doc, 'Subsector', level=3, style='Heading 2')
            # Año cerrado
            if (dict_verificacion['exportaciones_nme_cerrado'] == 'CON DATOS DE EXPORTACIONES NME CERRADO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CERRADO"]["SUBSECTORES"]), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            if (dict_verificacion['exportaciones_nme_corrido'] == 'CON DATOS DE EXPORTACIONES NME CORRIDO'):
                add_table(doc, pd.DataFrame(tablas["CATEGORIAS CORRIDO"]["SUBSECTORES"]), 'Table Grid', 10, fuente_exportaciones)

            ####################
            # Conteo de Empresas
            ####################
            if (dict_verificacion['exportaciones_conteo_cerrado'] == 'CON DATOS DE CONTEO CERRADO') or (dict_verificacion['exportaciones_conteo_corrido'] == 'CON DATOS DE CONTEO CORRIDO'):
                add_heading(doc, 'Empresas', level=3, style='Heading 2')
                # Año cerrado
                if (dict_verificacion['exportaciones_conteo_cerrado'] == 'CON DATOS DE CONTEO CERRADO'):
                    year_cerrado = doc_params['Año cerrado (T)']
                    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
                    conteo_cerrado = tablas['CONTEO EMPRESAS']['CERRADO'][0]
                    run_cerrado = paragraph_cerrado.add_run(f"{conteo_cerrado} empresas")
                    run_cerrado.bold = True
                elif (dict_verificacion['exportaciones_conteo_cerrado'] == 'SIN DATOS DE CONTEO CERRADO'):
                    year_cerrado = doc_params['Año cerrado (T)']
                    paragraph_cerrado = doc.add_paragraph(f'Número de empresas exportadoras en {str(year_cerrado)}: ', style='Normal')
                    conteo_cerrado = "0"
                    run_cerrado = paragraph_cerrado.add_run(f"{conteo_cerrado} empresas")
                    run_cerrado.bold = True                    
                # Año corrido
                if (dict_verificacion['exportaciones_conteo_corrido'] == 'CON DATOS DE CONTEO CORRIDO'):
                    year_corrido = doc_params['Año corrido texto (T)']
                    paragraph_corrido = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
                    conteo_corrido = tablas['CONTEO EMPRESAS']['CORRIDO'][0]
                    run_corrido = paragraph_corrido.add_run(f"{conteo_corrido} empresas")
                    run_corrido.bold = True
                elif (dict_verificacion['exportaciones_conteo_corrido'] == 'SIN DATOS DE CONTEO CORRIDO'):
                    year_corrido = doc_params['Año corrido texto (T)']
                    paragraph_corrido = doc.add_paragraph(f'Número de empresas exportadoras a {str(year_corrido)}: ', style='Normal')
                    conteo_corrido = "0"
                    run_corrido = paragraph_corrido.add_run(f"{conteo_corrido} empresas")
                    run_corrido.bold = True

            ################
            # Datos empresas
            ################
            if (dict_verificacion['exportaciones_empresas_cerrado'] == 'CON DATOS DE EMPRESAS CERRADO') or (dict_verificacion['exportaciones_empresas_corrido'] == 'CON DATOS DE EMPRESAS CORRIDO'):
                add_heading(doc, 'Información de Empresas', level=3, style='Heading 2')
                # Año cerrado
                if (dict_verificacion['exportaciones_empresas_cerrado'] == 'CON DATOS DE EMPRESAS CERRADO'):
                    add_table(doc, pd.DataFrame(tablas['EMPRESAS']['ST_NIT_CERRADO']), 'Table Grid', 9, fuente_exportaciones)
                # Año corrido
                if (dict_verificacion['exportaciones_empresas_corrido'] == 'CON DATOS DE EMPRESAS CORRIDO'):
                    add_table(doc, pd.DataFrame(tablas['EMPRESAS']['ST_NIT_CORRIDO']), 'Table Grid', 9, fuente_exportaciones)
            
            ##############################
            # Oportunidades de exportación
            ##############################
            if (dict_verificacion['oportunidades_exportacion'] == "CON OPORTUNIDADES"):
                add_heading(doc, 'Oportunidades de exportación identificadas', level=3, style='Heading 2')
                agregar_oportunidades_al_documento(doc, tablas, 'OPORTUNIDADES_EXPORTACIONES')
    
        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('El departamento no registra datos de exportaciones.')
    
    ###########
    # Inversión
    ###########
    if (dict_verificacion['oportunidades_inversion'] == "CON OPORTUNIDADES"):
            add_heading(doc, 'Inversión', level=2, style='Heading 1')
            add_heading(doc, 'Oportunidades de inversión identificadas', level=3, style='Heading 2')
            agregar_oportunidades_al_documento(doc, tablas, 'OPORTUNIDADES_INVERSION')

    #########
    # Turismo
    #########
    add_heading(doc, 'Turismo', level=2, style='Heading 1')
    if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO') or (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
        
        ########
        # Países
        ########
        add_heading(doc, 'Países', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["PAIS_RESIDENCIA"]), 'Table Grid', 10, fuente_turismo)
        # Año corrido
        if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CORRIDO"]["PAIS_RESIDENCIA"]), 'Table Grid', 10, fuente_turismo)

        ############
        # Municipios
        ############
        add_heading(doc, 'Municipios', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["CIUDAD_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)
        # Año corrido
        if (dict_verificacion['turismo_corrido'] == 'CON DATOS DE TURISMO CORRIDO'):
            add_table(doc, pd.DataFrame(tablas["TURISMO CORRIDO"]["CIUDAD_HOSPEDAJE"]), 'Table Grid', 10, fuente_turismo)

        ########
        # Género
        ########
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_heading(doc, 'Género', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["DESCRIPCION_GENERO"]), 'Table Grid', 10, fuente_turismo)

        #########
        # Motivos
        #########
        if (dict_verificacion['turismo_cerrado'] == 'CON DATOS DE TURISMO CERRADO'):
            add_heading(doc, 'Motivos', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas["TURISMO CERRADO"]["MOVC_NOMBRE"]), 'Table Grid', 10, fuente_turismo)
        
        ###############
        # OPORTUNIDADES
        ###############
        if (dict_verificacion['oportunidades_turismo'] == "CON OPORTUNIDADES"):
            add_heading(doc, 'Productos vacacionales que se promocionan', level=3, style='Heading 2')
            # Principales
            p1 = doc.add_paragraph()
            run1 = p1.add_run("Principales: ")
            run1.bold = True
            p1.add_run(tablas['TURISMO_PRINCIPAL'])
            p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # Nichos
            p2 = doc.add_paragraph()
            run2 = p2.add_run("Nichos: ")
            run2.bold = True
            p2.add_run(tablas['TURISMO_NICHOS'])
            p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Salto de página
        doc.add_page_break()
    else:
        doc.add_paragraph('El departamento no registra datos de turismo.')

    ##############
    # CONECTIVIDAD
    ##############
    if (dict_verificacion['conectividad'] == "CON DATOS DE CONECTIVIDAD"):
        add_heading(doc, 'Conectividad', level=2, style='Heading 1')
        add_table_resumen(doc, pd.DataFrame(tablas['CONECTIVIDAD']['CONECTIVIDAD']), 'Table Grid', 10, fuente_conectividad)
        # Salto de página
        doc.add_page_break()

    ###########
    # Logística
    ###########
    if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO') or (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
        add_heading(doc, 'Logística', level=2, style='Heading 1')
    #######
    # Pesos
    #######
        add_heading(doc, 'Pesos', level=3, style='Heading 2')
        # Año cerrado
        if (dict_verificacion['exportaciones_totales_cerrado'] == 'CON DATOS DE EXPORTACIONES TOTALES CERRADO'):
            add_table(doc, pd.DataFrame(tablas['TIPOS PESO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
        # Año corrido
        if (dict_verificacion['exportaciones_totales_corrido'] == 'CON DATOS DE EXPORTACIONES TOTALES CORRIDO'):
            add_table(doc, pd.DataFrame(tablas['TIPOS PESO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)
        # Mineros
        if (dict_verificacion['pesos_minero_cerrado'] == 'CON DATOS CERRADO') or (dict_verificacion['pesos_minero_corrido'] == 'CON DATOS CORRIDO'):
            add_heading(doc, 'Pesos por medio de transporte: exportaciones mineras', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO MINERO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO MINERO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)
        # No Mineros
        if (dict_verificacion['pesos_no_minero_cerrado'] == 'CON DATOS CERRADO') or (dict_verificacion['pesos_no_minero_corrido'] == 'CON DATOS CORRIDO'):
            add_heading(doc, 'Pesos por medio de transporte: exportaciones no mineras', level=3, style='Heading 2')
            # Año cerrado
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO NO MINERO']['ST_CATEGORIAS_PESO_CERRADO']), 'Table Grid', 10, fuente_exportaciones)
            # Año corrido
            add_table(doc, pd.DataFrame(tablas['MEDIOS PESO NO MINERO']['ST_CATEGORIAS_PESO_CORRIDO']), 'Table Grid', 10, fuente_exportaciones)

    ############
    # Disclaimer
    ############
    # Salto de página
    doc.add_page_break()
    # Agregar saltos de línea para centrar el texto verticalmente
    for _ in range(12): 
        doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
    # Agregar el texto del disclaimer
    paragraph_disclaimer = doc.add_paragraph(disclaimer)
    # Formato del texto
    run = paragraph_disclaimer.runs[0]
    run.font.name = 'Century Gothic'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 32, 96)
    run.bold = True
    paragraph_disclaimer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Guardar el documento
    doc.save(file_path)
